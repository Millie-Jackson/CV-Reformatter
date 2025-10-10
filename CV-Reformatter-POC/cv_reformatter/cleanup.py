from typing import Dict, Any
from docx.text.paragraph import Paragraph

def _norm(s: str) -> str:
    return " ".join((s or "").replace("\xa0", " ").strip().split()).upper()

def _remove_paragraph(p: Paragraph) -> None:
    p._element.getparent().remove(p._element)

def _remove_table_row_of(par: Paragraph) -> bool:
    el = par._element
    tc = el.getparent()
    if tc is None:
        return False
    tr = tc.getparent()
    if tr is None or not tr.tag.endswith('tr'):
        return False
    tbl = tr.getparent()
    try:
        tbl.remove(tr)
        return True
    except Exception:
        return False

def _row_text(row) -> str:
    try:
        return _norm("\n".join(c.text for c in row.cells))
    except Exception:
        return ""

def _looks_like_qual_guidance(text: str) -> bool:
    # Very specific, template-like phrases/headers (exact-ish matches)
    t = text.upper()
    return (
        "NAME OF ESTABLISHMENT" in t or
        "TITLE OF QUALIFICATION" in t or
        "DATE" == t or
        "LIST MOST RECENT FIRST." in t
    )

def prune_placeholders(doc, data: Dict[str, Any]) -> None:
    """
    Conservative cleanup:
    - If there's NO 'qualifications' data, remove the 'QUALIFICATIONS' heading
      only if we can identify the exact template element:
        * prefer removing the entire table row containing the heading text
        * otherwise remove the precise heading paragraph
      Optionally also remove the immediate next row/line if it is the guidance row.
    - Do NOT touch anything else.
    """
    if data.get("qualifications"):
        return  # user provided data -> keep the section

    # First try: if QUALIFICATIONS lives inside a table row, remove that row (and the next row if it's guidance)
    removed_any = False
    for tbl in doc.tables:
        for i, row in enumerate(list(tbl.rows)):
            rt = _row_text(row)
            if rt == "QUALIFICATIONS":
                try:
                    tbl._tbl.remove(row._tr)  # remove this row
                    removed_any = True
                    # If next row looks like the template guidance, remove it too (optional)
                    if i < len(tbl.rows):
                        next_rt = _row_text(tbl.rows[i])  # same index now points at the original next
                        if _looks_like_qual_guidance(next_rt):
                            tbl._tbl.remove(tbl.rows[i]._tr)
                except Exception:
                    pass
                break  # one match is enough
        if removed_any:
            break

    if removed_any:
        return

    # Fallback: remove a standalone QUALIFICATIONS heading paragraph (not in table)
    i = 0
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        if _norm(p.text) == "QUALIFICATIONS":
            _remove_paragraph(p)
            # If the immediately following paragraph is a guidance line, remove it too
            if i < len(doc.paragraphs):
                q = doc.paragraphs[i]
                if _looks_like_qual_guidance(_norm(q.text)):
                    _remove_paragraph(q)
            return  # done; do not keep scanning
        i += 1
