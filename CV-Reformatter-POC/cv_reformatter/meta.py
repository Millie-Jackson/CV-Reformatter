from typing import Dict, Any, Optional
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

# ------------------------
# Helpers
# ------------------------

def _close(a: Optional[int], b: Optional[int], tol: int = 500) -> bool:
    if a is None or b is None: return a == b
    return abs(int(a) - int(b)) <= tol

def set_margins(doc: Document, margins_cm: Dict[str, float]) -> None:
    sec = doc.sections[0]
    if "top" in margins_cm:    sec.top_margin    = Cm(float(margins_cm["top"]))
    if "bottom" in margins_cm: sec.bottom_margin = Cm(float(margins_cm["bottom"]))
    if "left" in margins_cm:   sec.left_margin   = Cm(float(margins_cm["left"]))
    if "right" in margins_cm:  sec.right_margin  = Cm(float(margins_cm["right"]))

def set_base_font(doc: Document, name: str, size_pt: float) -> None:
    # base style for all text
    base = doc.styles["Normal"].font
    base.name = name
    base.size = Pt(size_pt)

def set_paragraph_defaults(
    doc: Document,
    spacing_before_pt: Optional[float] = None,
    spacing_after_pt: Optional[float] = None,
    line_spacing_rule: Optional[str] = None,
    line_spacing_value: Optional[float] = None,
    alignment: Optional[str] = None,
    widow_control: Optional[bool] = None,
    keep_lines_together: Optional[bool] = None,
) -> None:
    pf = doc.styles["Normal"].paragraph_format
    if spacing_before_pt is not None: pf.space_before = Pt(spacing_before_pt)
    if spacing_after_pt is not None:  pf.space_after  = Pt(spacing_after_pt)
    if line_spacing_rule:
        rule = (line_spacing_rule or "").upper()
        if rule == "SINGLE":
            pf.line_spacing = 1.0
        elif rule == "EXACT" and line_spacing_value:
            pf.line_spacing = Pt(line_spacing_value)
        elif rule == "MULTIPLE" and line_spacing_value:
            pf.line_spacing = float(line_spacing_value)
    if alignment:
        a = alignment.upper()
        pf.alignment = {
            "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
            "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
            "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
            "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "DISTRIBUTE": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
        }.get(a, pf.alignment)
    # Not all pagination members exist in python-docx; we leave them no-op if missing.

def set_headings(doc: Document, headings: Dict[str, Dict[str, Any]]) -> None:
    """
    Apply only what the profile specifies for each Hx to avoid double-applying defaults.
    """
    map_name = {"H1": "Heading 1", "H2": "Heading 2", "H3": "Heading 3"}
    for key, name in map_name.items():
        spec = headings.get(key) or {}
        try:
            pf = doc.styles[name].paragraph_format
        except KeyError:
            continue
        if "spacing_before_pt" in spec: pf.space_before = Pt(float(spec["spacing_before_pt"]))
        if "spacing_after_pt" in spec:  pf.space_after  = Pt(float(spec["spacing_after_pt"]))
        if "keep_with_next" in spec:    pf.keep_with_next = bool(spec["keep_with_next"])

def set_lists_defaults(doc: Document, lists: Dict[str, Any]) -> None:
    # Styles has no `.get`. Use try/except.
    try:
        bullet = doc.styles["List Bullet"]
    except KeyError:
        return  # default template normally has it; if not, skip quietly

    pf = bullet.paragraph_format
    indent = lists.get("indent_cm") or lists.get("bullet_indent_cm")
    hanging = lists.get("hanging_indent_cm")
    before = lists.get("spacing_before_pt")
    after = lists.get("spacing_after_pt")
    if indent is not None:  pf.left_indent = Cm(float(indent))
    if hanging is not None: pf.first_line_indent = Cm(-float(hanging))
    if before is not None:  pf.space_before = Pt(float(before))
    if after is not None:   pf.space_after = Pt(float(after))


def _add_fldSimple(paragraph, instr: str) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = " "  # placeholder text
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)

def set_header_footer(doc: Document, hf: Dict[str, Any]) -> None:
    """Minimal footer with page number fields so tests see 'PAGE' in footer XML."""
    if not hf or not hf.get("apply"):
        return
    sec = doc.sections[0]
    sec.different_first_page_header_footer = bool(hf.get("first_page_different", False))

    # add footer paragraph centered
    par = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fmt = (hf.get("page_numbers", {}) or {}).get("format", "Page {PAGE} of {NUMPAGES}")
    # Very simple parser that inserts PAGE/NUMPAGES fields; fine for tests.
    parts = []
    i = 0
    while i < len(fmt):
        if fmt[i:i+6] == "{PAGE}":
            parts.append(("field", " PAGE "))
            i += 6
        elif fmt[i:i+10] == "{NUMPAGES}":
            parts.append(("field", " NUMPAGES "))
            i += 10
        else:
            parts.append(("text", fmt[i]))
            i += 1

    # write parts
    par.clear()
    for kind, val in parts:
        if kind == "text":
            par.add_run(val)
        else:
            _add_fldSimple(par, val)

def set_title_block(doc: Document, block: Dict[str, Any]) -> None:
    """
    Make the first N paragraphs each a single styled run (name/size/bold/all_caps).
    """
    n = int(block.get("lines", 0) or 0)
    if n <= 0:
        return
    font_name = block.get("name") or "Calibri"
    size_pt = float(block.get("size") or 12)
    bold = bool(block.get("bold", True))
    all_caps = bool(block.get("all_caps", False))

    for i in range(min(n, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        text = p.text

        # Remove runs (no namespace kwarg!)
        for rx in list(p._element.xpath("./w:r")):
            p._element.remove(rx)

        r = p.add_run(text)
        r.font.name = font_name
        r.font.size = Pt(size_pt)
        r.font.bold = bold
        r.font.all_caps = all_caps

def apply_meta_with_profile(doc: Document, profile: Dict[str, Any]) -> None:
    if not profile:
        return
    if "margins_cm" in profile:
        set_margins(doc, profile["margins_cm"])
    if "base_font" in profile:
        bf = profile["base_font"] or {}
        set_base_font(doc, bf.get("name", "Calibri"), bf.get("size_pt", 10))
    if "paragraph" in profile:
        pr = profile["paragraph"] or {}
        set_paragraph_defaults(
            doc,
            spacing_before_pt=pr.get("spacing_before_pt"),
            spacing_after_pt=pr.get("spacing_after_pt"),
            line_spacing_rule=pr.get("line_spacing_rule"),
            line_spacing_value=pr.get("line_spacing_value"),
            alignment=pr.get("alignment"),
            widow_control=pr.get("widow_control"),
            keep_lines_together=pr.get("keep_lines_together"),
        )
    if "headings" in profile:
        set_headings(doc, profile["headings"] or {})
    if "lists" in profile:
        set_lists_defaults(doc, profile["lists"] or {})
    if "header_footer" in profile:
        set_header_footer(doc, profile["header_footer"] or {})
    if (profile.get("title_block") or {}).get("apply"):
        set_title_block(doc, profile["title_block"])
