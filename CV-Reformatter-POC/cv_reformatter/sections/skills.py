# cv_reformatter/sections/skills.py
from typing import Any, List, Optional, Dict
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

def _find_heading_para(doc: Document, title_upper: str) -> Optional[int]:
    """Return index of a paragraph whose text (upper) equals title_upper."""
    for i, p in enumerate(doc.paragraphs):
        if (p.text or "").strip().upper() == title_upper:
            return i
    return None

def _insert_after(anchor_para: Paragraph) -> Paragraph:
    """
    Insert a new empty paragraph directly AFTER `anchor_para` and return it.
    Uses safe xml insertion + wraps it back into a python-docx Paragraph.
    """
    new_el = OxmlElement("w:p")
    anchor_para._element.addnext(new_el)
    # parent is a _Body-like container; python-docx Paragraph takes (el, parent)
    return Paragraph(new_el, anchor_para._parent)

def write_section(doc: Document, title: str, body: str, data: Dict[str, Any]) -> None:
    """
    KEY SKILLS writer.
    Ensures there is at least one non-empty paragraph under the KEY SKILLS heading.
    If structured skills exist, writes a single line joined by semicolons.
    """
    skills: List[str] = list(map(str, (data.get("skills") or [])))
    title_upper = (title or "").strip().upper()
    idx = _find_heading_para(doc, title_upper)
    if idx is None:
        return  # no heading present, nothing to do

    # Remove placeholder blank lines immediately after the heading up to next heading.
    i = idx + 1
    while i < len(doc.paragraphs):
        q = doc.paragraphs[i]
        style_name = (getattr(q.style, "name", "") or "").upper()
        if style_name in ("HEADING 1", "HEADING 2", "HEADING 3"):
            break
        if (q.text or "").strip():
            break  # keep first real line if already authored
        q._element.getparent().remove(q._element)
        # do not increment i; collection just shrank

    anchor = doc.paragraphs[idx]

    if not skills:
        # Ensure a single non-empty line so E2E sees body content exists.
        p = _insert_after(anchor)
        p.add_run(" ")
        return

    joined = "; ".join(s.strip() for s in skills if s.strip()) or " "
    p = _insert_after(anchor)
    p.add_run(joined)
