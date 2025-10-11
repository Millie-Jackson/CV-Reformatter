from __future__ import annotations
from typing import Any, Dict, Optional, Tuple
import re

from docx.document import Document as _Doc  # type: ignore
from docx.table import Table  # type: ignore
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

# ----- constants -----
HEADER_FONT_NAME = "Calibri"
HEADER_FONT_SIZE_PT = 12
_ZWS = "\u200B"

# ----- tiny helpers -----
def _norm(s: Optional[str]) -> str:
    return (s or "").strip()

def _upper(s: Optional[str]) -> str:
    return _norm(s).upper()

def _first(*vals: Optional[str]) -> Optional[str]:
    for v in vals:
        if _norm(v):
            return _norm(v)
    return None

def _delete_paragraph(p: Paragraph) -> None:
    elm = p._element
    parent = elm.getparent()
    if parent is not None:
        parent.remove(elm)

def _force_normal_style(p: Paragraph) -> None:
    try:
        p.style = p.part.document.styles["Normal"]
    except Exception:
        try:
            p.style = "Normal"
        except Exception:
            pass

def _center_bold12(p: Paragraph) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.keep_with_next = True
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.bold = True
        r.font.size = Pt(HEADER_FONT_SIZE_PT)
        r.font.name = HEADER_FONT_NAME
        try:
            rpr = r._r.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:ascii"),   HEADER_FONT_NAME)
            rfonts.set(qn("w:hAnsi"),   HEADER_FONT_NAME)
            rfonts.set(qn("w:eastAsia"),HEADER_FONT_NAME)
            rfonts.set(qn("w:cs"),      HEADER_FONT_NAME)
        except Exception:
            pass

def _set_text(p: Paragraph, text: str) -> None:
    _force_normal_style(p)
    for r in list(p.runs):
        r.clear()
        p._element.remove(r._r)
    p.add_run(text)
    _center_bold12(p)

def _style_spacer(p: Paragraph) -> None:
    _force_normal_style(p)
    for r in list(p.runs):
        r.clear()
        p._element.remove(r._r)
    p.add_run(_ZWS)
    _center_bold12(p)

def _iter_pars_in_table(t: Table):
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for t2 in cell.tables:
                yield from _iter_pars_in_table(t2)

def _iter_all_pars(doc: _Doc):
    for p in doc.paragraphs:
        yield p
    for s in doc.sections:
        for p in s.header.paragraphs:
            yield p
        for t in s.header.tables:
            yield from _iter_pars_in_table(t)
        for p in s.footer.paragraphs:
            yield p
        for t in s.footer.tables:
            yield from _iter_pars_in_table(t)

def _has_border(p: Paragraph) -> bool:
    try:
        return bool(p._element.xpath(".//w:pPr/w:pBdr", namespaces=p._element.nsmap))
    except Exception:
        return False

def _has_drawing(p: Paragraph) -> bool:
    try:
        return bool(p._element.xpath(".//w:drawing", namespaces=p._element.nsmap))
    except Exception:
        return False

def _is_blank(p: Paragraph) -> bool:
    txt = p.text or ""
    for ch in ("\u00A0", "\u2000", "\u2001", "\u2002", "\u2003", "\u2004",
               "\u2005", "\u2006", "\u2007", "\u2008", "\u2009", "\u200A",
               "\u200B", "\u202F", "\u205F", "\u3000"):
        txt = txt.replace(ch, " ")
    return txt.strip() == ""

def _insert_before(doc: _Doc, ref_idx: int, new_p: Paragraph) -> None:
    dest = doc.paragraphs[ref_idx]._element if ref_idx < len(doc.paragraphs) else None
    elm = new_p._element
    body = doc._body._body
    body.remove(elm)
    if dest is None:
        body.append(elm)
    else:
        body.insert(body.index(dest), elm)

# ----- fallback extraction from existing text -----
_name_re = re.compile(r"^CURRICULUM\s+VITAE\s+FOR\s+(?P<n>.+?)\s*$", re.I)
_loc_re  = re.compile(r"^CANDIDATE\s+LOCATION:\s+(?P<l>.+?)\s*$", re.I)

def _extract_from_doc(doc: _Doc) -> Tuple[Optional[str], Optional[str]]:
    name = None
    loc  = None
    for p in _iter_all_pars(doc):
        t = (p.text or "").strip()
        if not name:
            m = _name_re.match(t)
            if m: name = m.group("n").strip()
        if not loc:
            m = _loc_re.match(t)
            if m: loc = m.group("l").strip()
        if name and loc:
            break
    return name, loc

# ----- main header routine -----
def write_section(doc: _Doc, title: str, body: str, data: Dict[str, Any]) -> None:
    # Resolve name/location – prefer structured data; fall back to doc; then safe placeholders
    hdr = data.get("header") if isinstance(data.get("header"), dict) else {}
    name = _first(
        hdr.get("name") if isinstance(hdr, dict) else None,
        data.get("candidate_name"), data.get("name"),
        " ".join([_norm(data.get("first_name") or ""), _norm(data.get("last_name") or "")]).strip() or None
    )
    loc  = _first(
        hdr.get("location") if isinstance(hdr, dict) else None,
        data.get("candidate_location"), data.get("location")
    )
    if not name or not loc:
        n2, l2 = _extract_from_doc(doc)
        name = name or n2
        loc  = loc  or l2
    name_u = _upper(name) or "FIRSTNAME LASTNAME"
    loc_u  = _upper(loc)   if loc else None

    # Remove template placeholders if they exist anywhere
    for p in list(_iter_all_pars(doc)):
        t = _upper(p.text)
        if t in {
            "CURRICULUM VITAE FOR FIRSTNAME LASTNAME",
            "CANDIDATE LOCATION: N/A"
        }:
            _delete_paragraph(p)

    # 1) Find the logo
    logo_idx = None
    for i, p in enumerate(doc.paragraphs):
        if _has_drawing(p):
            logo_idx = i
            break
    if logo_idx is None:
        # No logo – just ensure there is exactly one pair of title lines at the top
        _ensure_titles_somewhere(doc, name_u, loc_u)
        return

    # 2) Ensure EXACTLY ONE spacer under the logo
    j = logo_idx + 1
    if j < len(doc.paragraphs) and not _has_border(doc.paragraphs[j]) and _is_blank(doc.paragraphs[j]):
        _style_spacer(doc.paragraphs[j])
        spacer_idx = j
    else:
        sp = doc.add_paragraph("")
        _style_spacer(sp)
        _insert_before(doc, j, sp)
        spacer_idx = j

    # 3) Find separator (first bordered paragraph) after spacer
    k = spacer_idx + 1
    while k < len(doc.paragraphs) and not _has_border(doc.paragraphs[k]):
        k += 1  # k points to separator or end

    # 4) Between spacer and separator:
    #    - remove only pure-blank lines (keep any existing titles)
    #    - track if lines already exist
    have_name = False
    have_loc  = False
    scan = spacer_idx + 1
    while scan < k and scan < len(doc.paragraphs):
        t = _upper(doc.paragraphs[scan].text)
        if not t and _is_blank(doc.paragraphs[scan]):
            _delete_paragraph(doc.paragraphs[scan])
            k -= 1
            continue
        if t.startswith("CURRICULUM VITAE FOR"):
            have_name = True
            _set_text(doc.paragraphs[scan], f"CURRICULUM VITAE FOR {name_u}".rstrip())
        elif t.startswith("CANDIDATE LOCATION:"):
            have_loc = True
            if loc_u:
                _set_text(doc.paragraphs[scan], f"CANDIDATE LOCATION: {loc_u}")
            else:
                _delete_paragraph(doc.paragraphs[scan]); k -= 1
                have_loc = False
                scan -= 1
        scan += 1

    # 5) Insert missing lines just before separator (or end)
    if not have_name:
        p_name = doc.add_paragraph("")
        _set_text(p_name, f"CURRICULUM VITAE FOR {name_u}".rstrip())
        _insert_before(doc, k, p_name)
        k += 1
    if loc_u and not have_loc:
        p_loc = doc.add_paragraph("")
        _set_text(p_loc, f"CANDIDATE LOCATION: {loc_u}")
        _insert_before(doc, k, p_loc)
        k += 1

    # 6) Make separator tight; preserve following heading spacing
    if k < len(doc.paragraphs) and _has_border(doc.paragraphs[k]):
        doc.paragraphs[k].paragraph_format.space_before = Pt(0)
        nxt = k + 1
    else:
        nxt = k
    if nxt < len(doc.paragraphs):
        try:
            pf = doc.styles["Heading 2"].paragraph_format
            target = pf.space_before
        except Exception:
            target = Pt(0)
        doc.paragraphs[nxt].paragraph_format.space_before = target if target is not None else Pt(0)

def _ensure_titles_somewhere(doc: _Doc, name_u: str, loc_u: Optional[str]) -> None:
    """
    Fallback when no logo is present: ensure the two lines exist near
    the top of the document (in Normal style, centered, bold 12).
    """
    inserted_name = inserted_loc = False
    # try update in-place if found
    for p in doc.paragraphs[:10]:
        t = _upper(p.text)
        if t.startswith("CURRICULUM VITAE FOR"):
            _set_text(p, f"CURRICULUM VITAE FOR {name_u}".rstrip())
            inserted_name = True
        elif t.startswith("CANDIDATE LOCATION:"):
            if loc_u:
                _set_text(p, f"CANDIDATE LOCATION: {loc_u}")
                inserted_loc = True
            else:
                _delete_paragraph(p)
        if inserted_name and (inserted_loc or not loc_u):
            return

    # otherwise add them at the very top
    p0 = doc.add_paragraph("")
    _set_text(p0, f"CURRICULUM VITAE FOR {name_u}".rstrip())
    doc._body._body.insert(0, p0._element)
    if loc_u:
        p1 = doc.add_paragraph("")
        _set_text(p1, f"CANDIDATE LOCATION: {loc_u}")
        doc._body._body.insert(1, p1._element)
