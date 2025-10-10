from typing import List, Optional, Iterable, Any
from docx.text.paragraph import Paragraph

TARGET_LABELS = {"OTHER HEADINGS", "OTHER HEADING", "PROFESSIONAL DEVELOPMENT"}

def _norm(s: str) -> str:
    return " ".join((s or "").replace("\xa0", " ").strip().split()).upper()

def _in_targets(s: str) -> bool:
    return _norm(s) in TARGET_LABELS

def _row_of_paragraph(p: Paragraph):
    el = p._element
    tc = el.getparent()
    if tc is None or not tc.tag.endswith("tc"):
        return None
    tr = tc.getparent()
    if tr is None or not tr.tag.endswith("tr"):
        return None
    tbl = tr.getparent()
    if tbl is None or not tbl.tag.endswith("tbl"):
        return None
    return tbl, tr  # raw oxml

def _find_row_any_cell_text(doc) -> Optional[tuple]:
    # 1) fast path: any cell in any row matches our target labels
    for tbl in doc.tables:
        for i, row in enumerate(tbl.rows):
            try:
                cell_texts = [_norm(c.text) for c in row.cells]
            except Exception:
                continue
            if any(_in_targets(t) or t == "OTHER HEADINGS" for t in cell_texts):
                return (tbl, i)
    # 2) paragraph-level: if label is a paragraph inside a cell
    for p in doc.paragraphs:
        if not _in_targets(p.text):
            continue
        mapped = _row_of_paragraph(p)
        if not mapped:
            continue
        raw_tbl, raw_tr = mapped
        for tbl in doc.tables:
            if tbl._tbl is raw_tbl:
                for i, row in enumerate(tbl.rows):
                    if row._tr is raw_tr:
                        return (tbl, i)
    return None

def _clear_cell(cell) -> None:
    cell.text = ""

def _add_bullets(cell, items: List[str]) -> None:
    for item in items:
        p = cell.add_paragraph(item)
        try:
            p.style = "List Bullet"
        except Exception:
            pass

def _coerce_items(items: Any) -> List[str]:
    """Accept list[str] or list[dict] (use common text fields)."""
    if not isinstance(items, list):
        return []
    out: List[str] = []
    for it in items:
        if isinstance(it, str):
            s = it.strip()
            if s:
                out.append(s)
        elif isinstance(it, dict):
            for k in ("text", "name", "title", "value", "label"):
                v = it.get(k)
                if isinstance(v, str) and v.strip():
                    out.append(v.strip())
                    break
    return out

def write_professional_development(doc, items_any) -> None:
    # Always try to find and rename the row so you see the section label change
    hit = _find_row_any_cell_text(doc)
    items = _coerce_items(items_any)

    if hit:
        tbl, idx = hit
        row = tbl.rows[idx]
        # Rename whichever cell contains the label to "PROFESSIONAL DEVELOPMENT"
        for c in row.cells:
            if _in_targets(c.text):
                c.text = "PROFESSIONAL DEVELOPMENT"
                break
        # If we have items, populate the right-most cell as bullets
        if items:
            content_cell = row.cells[-1]
            _clear_cell(content_cell)
            _add_bullets(content_cell, items)
        return

    # Fallback: append to end of doc (rare)
    h = doc.add_paragraph("PROFESSIONAL DEVELOPMENT")
    try:
        if "Heading 2" in doc.styles:
            h.style = doc.styles["Heading 2"]
        elif "Heading 1" in doc.styles:
            h.style = doc.styles["Heading 1"]
    except Exception:
        pass
    for it in items:
        p = doc.add_paragraph(it)
        try:
            p.style = "List Bullet"
        except Exception:
            pass
