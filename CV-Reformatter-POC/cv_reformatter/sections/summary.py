from __future__ import annotations
from typing import Any, Dict, Optional

from docx.document import Document as _Doc  # type: ignore
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

# ---------- small helpers ----------
def _norm(s: Optional[str]) -> str:
    return (s or "").strip()

def _upper(s: Optional[str]) -> str:
    return _norm(s).upper()

def _has_paragraph_border(p: Paragraph) -> bool:
    try:
        return bool(p._element.xpath(".//w:pPr/w:pBdr", namespaces=p._element.nsmap))
    except Exception:
        return False

def _delete_paragraph(p: Paragraph) -> None:
    elm = p._element
    parent = elm.getparent()
    if parent is not None:
        parent.remove(elm)

def _insert_before_ref_paragraph(ref_p: Paragraph, new_p: Paragraph) -> None:
    """Insert new_p immediately before ref_p, regardless of parent container."""
    new_el = new_p._element
    par = new_el.getparent()
    if par is not None:
        par.remove(new_el)

    ref_el = ref_p._element
    parent = ref_el.getparent()
    if parent is None:
        body = ref_p.part.document._body._body
        body.append(new_el)
        return

    parent.insert(parent.index(ref_el), new_el)

def _justify_normal(p: Paragraph) -> None:
    pf = p.paragraph_format
    pf.space_before = None
    pf.space_after = None
    pf.keep_together = False
    pf.keep_with_next = False
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def _looks_like_placeholder(text: str) -> bool:
    t = (text or "").strip().lower()
    # be tolerant to tiny variations
    return (
        t.startswith("<insert") and
        ("executive" in t) and
        ("summary" in t) and
        t.endswith(">")
    )

# ---------- public ----------
def write_section(doc: _Doc, title: str, body: str, data: Dict[str, Any]) -> None:
    """
    Replace the PERSONAL PROFILE placeholder with the extracted summary,
    placing it between the two bordered paragraphs that frame the section.
    Also remove any remaining placeholder line(s).
    """
    # 1) best available summary text
    src = data or {}
    summary = (
        src.get("personal_profile")
        or src.get("summary")
        or src.get("executive_summary")
        or src.get("profile")
        or ""
    )
    summary = _norm(summary)
    if not summary:
        return

    # 2) locate heading
    h_idx = None
    for i, p in enumerate(doc.paragraphs):
        if _upper(p.text).startswith("PERSONAL PROFILE"):
            h_idx = i
            break
    if h_idx is None:
        return

    # 3) top rule after heading
    b1 = h_idx + 1
    while b1 < len(doc.paragraphs) and not _has_paragraph_border(doc.paragraphs[b1]):
        b1 += 1

    # If no rule at all, insert right after the heading and remove any placeholders nearby
    if b1 >= len(doc.paragraphs):
        newp = doc.add_paragraph("")
        newp.add_run(summary)
        _justify_normal(newp)
        # insert after heading
        ref_el = doc.paragraphs[h_idx]._element
        parent = ref_el.getparent()
        pel = newp._element
        par = pel.getparent()
        if par is not None:
            par.remove(pel)
        if parent is not None:
            parent.insert(parent.index(ref_el) + 1, pel)
        # cleanup placeholders until the next heading or border
        i = h_idx + 1
        while i < len(doc.paragraphs):
            q = doc.paragraphs[i]
            if _has_paragraph_border(q):
                break
            if _upper(q.text) in ("KEY SKILLS","EDUCATION","EMPLOYMENT HISTORY",
                                  "PROFESSIONAL DEVELOPMENT","ADDITIONAL INFORMATION","OTHER HEADINGS"):
                break
            if _looks_like_placeholder(q.text):
                _delete_paragraph(q)
                continue
            i += 1
        return

    # 4) bottom rule (next border after b1)
    b2 = b1 + 1
    while b2 < len(doc.paragraphs) and not _has_paragraph_border(doc.paragraphs[b2]):
        b2 += 1
    if b2 > len(doc.paragraphs):
        b2 = len(doc.paragraphs)

    # 5) delete everything strictly between b1 and b2 (keep the two rules)
    idx = b1 + 1
    while idx < b2 and idx < len(doc.paragraphs):
        _delete_paragraph(doc.paragraphs[idx])
        b2 -= 1  # account for shrinking list

    # 6) insert summary JUST BEFORE bottom rule (or append in same container if no bottom rule)
    newp = doc.add_paragraph("")
    newp.add_run(summary)
    _justify_normal(newp)

    if b2 < len(doc.paragraphs):
        _insert_before_ref_paragraph(doc.paragraphs[b2], newp)
    else:
        # append to the same container as the top rule
        top_rule_el = doc.paragraphs[b1]._element
        parent = top_rule_el.getparent()
        pel = newp._element
        par = pel.getparent()
        if par is not None:
            par.remove(pel)
        if parent is not None:
            parent.append(pel)
        else:
            doc._body._body.append(pel)

    # 7) defensive cleanup: remove any placeholder lines remaining between the two rules
    #    (rare template variants where borders are mis-detected).
    # Recompute bounds (doc.paragraphs changed after insert).
    # Find b1 again
    start = None
    end = None
    for i, p in enumerate(doc.paragraphs):
        if i > h_idx and _has_paragraph_border(p):
            start = i
            break
    if start is not None:
        j = start + 1
        while j < len(doc.paragraphs) and not _has_paragraph_border(doc.paragraphs[j]):
            j += 1
        end = j if j < len(doc.paragraphs) else len(doc.paragraphs)

    if start is not None and end is not None:
        i = start + 1
        while i < end and i < len(doc.paragraphs):
            q = doc.paragraphs[i]
            if _looks_like_placeholder(q.text):
                _delete_paragraph(q)
                end -= 1
                continue
            i += 1
