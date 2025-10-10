# cv_reformatter/sections/education.py
from __future__ import annotations
from typing import Any, Dict, List, Optional
from docx import Document
from docx.text.paragraph import Paragraph

# Accept common keys users might have in fields.json
_EDU_KEYS = ["education", "qualifications", "training"]

def _s(x: Any) -> Optional[str]:
    if x is None:
        return None
    xs = str(x).strip()
    return xs if xs else None

def _get_education_list(data: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Normalize education entries into a list of:
       {year, institution, title, award, bullets: List[str]}
       Keep 'most recent first' when a numeric year is present; otherwise preserve input order.
    """
    src: Optional[Any] = None
    for k in _EDU_KEYS:
        if k in data and data[k] not in (None, "", []):
            src = data[k]
            break
    if src is None:
        return []

    # normalize to list
    entries: List[Dict[str, Any]] = []
    if isinstance(src, list):
        it = src
    elif isinstance(src, dict):
        it = [src]
    else:
        it = [{"title": str(src)}]

    for raw in it:
        entries.append(_normalize_entry(raw))

    # sort by numeric 'year' desc when possible; otherwise stable
    def _year_num(e: Dict[str, Any]) -> int:
        y = e.get("year")
        if not y:
            return -10**9  # push unknown years to end while preserving original order
        try:
            # handle '2021', '2018-2020', etc. -> take first 4-digit number
            import re
            m = re.search(r"\d{4}", str(y))
            return int(m.group(0)) if m else -10**9
        except Exception:
            return -10**9

    # Keep stable order for equal keys by using enumerate index
    entries_with_idx = list(enumerate(entries))
    entries_with_idx.sort(key=lambda t: (_year_num(t[1]), t[0]), reverse=True)
    return [e for _, e in entries_with_idx]

def _normalize_entry(raw: Any) -> Dict[str, Any]:
    if not isinstance(raw, dict):
        return {
            "year": None,
            "institution": None,
            "title": _s(raw),
            "award": None,
            "bullets": [],
        }

    def g(*keys, default=None):
        for k in keys:
            if k in raw and raw[k] not in (None, "", []):
                return raw[k]
        return default

    year        = g("year", "date", "when", "graduation_year")
    institution = g("institution", "establishment", "school", "university", "college", "name")
    title       = g("title", "qualification", "degree", "program", "course")
    award       = g("award", "result", "grade", "honours", "honors")
    bullets     = g("bullets", "modules", "highlights", default=[])

    # normalize bullets as list[str]
    if isinstance(bullets, str):
        bullets = [b.strip() for b in bullets.split("\n") if b.strip()]
    elif isinstance(bullets, list):
        bullets = [str(b).strip() for b in bullets if str(b).strip()]
    else:
        bullets = []

    return {
        "year": _s(year),
        "institution": _s(institution),
        "title": _s(title),
        "award": _s(award),
        "bullets": bullets,
    }

def _add_heading(doc: Document, text: str) -> Paragraph:
    hp = doc.add_paragraph(text)
    try:
        hp.style = "Heading 2"
    except Exception:
        pass
    return hp

def _join_nonempty(parts: List[Optional[str]], sep: str) -> str:
    vals = [p for p in parts if isinstance(p, str) and p.strip()]
    return sep.join(vals)

def write_section(doc: Document, title: str, body: str, data: Dict[str, Any]) -> None:
    """
    Writer used by pipeline dispatch.
    Renders:
      - Heading "EDUCATION" (H2)
      - For each entry (most recent first, if year available):
         Line 1: Year \t Institution
         Line 2: Title (and Award if present: 'Title, Award')
         Bullets (optional): List Bullet
    """
    _add_heading(doc, "EDUCATION")
    entries = _get_education_list(data)
    if not entries:
        return

    for e in entries:
        # Line 1
        line1 = _join_nonempty([e.get("year"), e.get("institution")], "\t")
        doc.add_paragraph(line1)

        # Line 2
        t = e.get("title")
        a = e.get("award")
        if t or a:
            line2 = _join_nonempty(
                [t, a] if (t and a) else ([t] if t else [a]),
                ", "
            )
            doc.add_paragraph(line2)

        # Optional bullets
        for b in e.get("bullets") or []:
            p = doc.add_paragraph()
            try:
                p.style = "List Bullet"
            except Exception:
                pass
            p.add_run(b)
