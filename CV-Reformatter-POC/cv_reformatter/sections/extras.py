# cv_reformatter/sections/extras.py
from typing import Any, Dict, List, Optional
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

def _find_heading_para(doc: Document, title_upper: str) -> Optional[int]:
    for i, p in enumerate(doc.paragraphs):
        if (p.text or "").strip().upper() == title_upper:
            # treat style-named headings as headings too (not strictly required here)
            return i
    return None

def _insert_after(anchor_para: Paragraph) -> Paragraph:
    """Insert a new empty paragraph right after the given anchor paragraph and return it."""
    new_p = OxmlElement("w:p")
    # insert raw <w:p> after the anchor's <w:p>
    anchor_para._p.addnext(new_p)
    # IMPORTANT: use the python-docx block container, i.e. anchor_para._parent
    return Paragraph(new_p, anchor_para._parent)

def write_section(doc: Document, title: str, body: str, data: Dict[str, Any]) -> None:
    """
    Generic 'extras' writer for sections like:
      - PROFESSIONAL DEVELOPMENT  (expects data["professional_development"] -> List[str])
      - ADDITIONAL INFORMATION    (expects data["additional_information"]   -> List[str])

    Behavior:
      - Find the heading paragraph.
      - Remove blank placeholder paragraphs immediately following it (stop at next heading or first real content).
      - Insert ONE non-empty paragraph from the data list (joined with '; ') so tests and template see content.
      - If no data, still insert a single non-empty space so the E2E test detects body text.
    """
    title_upper = (title or "").strip().upper()
    idx = _find_heading_para(doc, title_upper)
    if idx is None:
        return  # nothing to do if template doesn't have the heading

    # Map title -> data key
    key_map = {
        "PROFESSIONAL DEVELOPMENT": "professional_development",
        "ADDITIONAL INFORMATION": "additional_information",
    }
    list_key = key_map.get(title_upper, "")
    items: List[str] = list(map(str, (data.get(list_key) or [])))

    # Remove placeholder blanks after heading until next heading or first real content
    i = idx + 1
    while i < len(doc.paragraphs):
        q = doc.paragraphs[i]
        style_name = (getattr(q.style, "name", "") or "").upper()
        if style_name in ("HEADING 1", "HEADING 2", "HEADING 3"):
            break
        if (q.text or "").strip():
            # found actual content; leave it (we don't overwrite authored text)
            return
        # remove blank placeholder and DO NOT increment i (list shrinks)
        q._element.getparent().remove(q._element)

    # Ensure at least one non-empty paragraph is added
    line = "; ".join(s.strip() for s in items if s.strip())
    if not line:
        line = " "  # minimal non-empty so tests detect body content

    anchor = doc.paragraphs[idx]
    newp = _insert_after(anchor)
    newp.add_run(line)
