# cv_reformatter/sections/experience.py
from __future__ import annotations
from typing import Any, Dict, List, Optional
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt


# ---------------------------
# Helpers: extraction/normalization
# ---------------------------
_KEYS = ["employment_history", "experience", "work_history"]

def _get_experience_list(data: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Return a list of role dicts with a tolerant schema:

    Each role will be normalized to:
      {
        "title": str|None,
        "company": str|None,
        "location": str|None,
        "start": str|None,        # 'Start Date'
        "end": str|None,          # 'End Date' or 'Present'
        "company_blurb": str|None,
        "bullets": List[str]      # may be empty
      }
    """
    src: Optional[Any] = None
    for k in _KEYS:
        if k in data and data[k]:
            src = data[k]
            break

    if src is None:
        return []

    roles: List[Dict[str, Any]] = []
    if isinstance(src, list):
        for r in src:
            roles.append(_normalize_role(r))
    elif isinstance(src, dict):
        # single role as dict
        roles.append(_normalize_role(src))
    else:
        # best-effort wrap
        roles.append(_normalize_role({"title": str(src)}))
    # most recent first if dates exist (very light — keep stable if missing)
    def _sort_key(r: Dict[str, Any]):
        # we don't parse dates; we just bias to keep original order
        return 0
    return sorted(roles, key=_sort_key)


def _normalize_role(raw: Any) -> Dict[str, Any]:
    if not isinstance(raw, dict):
        return {
            "title": str(raw).strip() if raw else None,
            "company": None,
            "location": None,
            "start": None,
            "end": None,
            "company_blurb": None,
            "bullets": [],
        }
    # Allow multiple common keys
    def g(*keys, default=None):
        for k in keys:
            v = raw.get(k)
            if v not in (None, "", []):
                return v
        return default

    title   = g("title", "role", "job_title")
    company = g("company", "employer", "org", "organisation", "organization")
    location= g("location", "city", "place")
    start   = g("start", "start_date", "from", "date_from")
    end     = g("end", "end_date", "to", "date_to", default="Present")
    blurb   = g("company_blurb", "company_info", "employer_summary", "about_company")
    bullets = g("bullets", "responsibilities", "highlights", default=[])

    # normalize bullets to list[str]
    if isinstance(bullets, str):
        bullets = [b.strip() for b in bullets.split("\n") if b.strip()]
    elif isinstance(bullets, list):
        bullets = [str(b).strip() for b in bullets if str(b).strip()]
    else:
        bullets = []

    return {
        "title": _s(title),
        "company": _s(company),
        "location": _s(location),
        "start": _s(start),
        "end": _s(end),
        "company_blurb": _s(blurb),
        "bullets": bullets,
    }


def _s(x: Any) -> Optional[str]:
    if x is None:
        return None
    xs = str(x).strip()
    return xs if xs else None


# ---------------------------
# Writer API
# ---------------------------
def write_section(doc: Document, title: str, body: str, data: Dict[str, Any]) -> None:
    """
    Public entrypoint used by pipeline dispatch.
    - Adds 'EMPLOYMENT HISTORY' heading (H2)
    - Renders role blocks:
        Line 1: Start – End \t COMPANY, Location
        Line 2: Job Title  (bold)
        Line 3: Company blurb (italic, optional)
        Bullets: List Bullet
    Notes:
      - Idempotency at the pipeline level: writer is called once per build on a fresh doc.
      - We lean on global meta for list spacing/indents.
    """
    # Heading
    _add_heading(doc, "EMPLOYMENT HISTORY")

    roles = _get_experience_list(data)
    if not roles:
        # If no roles, keep the heading to preserve section structure
        return

    for r in roles:
        _write_role_block(doc, r)


def _add_heading(doc: Document, text: str) -> Paragraph:
    hp = doc.add_paragraph(text)
    try:
        hp.style = "Heading 2"
    except Exception:
        pass
    return hp


def _write_role_block(doc: Document, r: Dict[str, Any]) -> None:
    # Header line: dates + company/location
    dates = _join_nonempty([r.get("start"), "–", r.get("end")], sep=" ")
    org   = _join_nonempty([r.get("company"), r.get("location")], sep=", ")
    header_line = _join_nonempty([dates, org], sep="\t")
    p_hdr = doc.add_paragraph(header_line)

    # Role title line (bold)
    title = r.get("title")
    if title:
        p_title = doc.add_paragraph()
        run = p_title.add_run(title)
        run.bold = True

    # Company blurb (italic)
    blurb = r.get("company_blurb")
    if blurb:
        p_blurb = doc.add_paragraph()
        run = p_blurb.add_run(blurb)
        run.italic = True

    # Bullets
    for b in r.get("bullets") or []:
        p = doc.add_paragraph()
        try:
            p.style = "List Bullet"
        except Exception:
            pass
        p.add_run(b)


def _join_nonempty(items: List[Optional[str]], sep: str) -> str:
    vals = [x for x in items if isinstance(x, str) and x.strip()]
    return sep.join(vals)
