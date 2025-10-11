import json
import os
from docx import Document
from docx.shared import Pt
from cv_reformatter.pipeline import reformat_cv_cv1_to_template1

def _make_minimal_template(path: str) -> None:
    doc = Document()
    # headings only; writers will populate bodies
    for title in [
        "PERSONAL PROFILE",
        "KEY SKILLS",
        "PROFESSIONAL DEVELOPMENT",
        "EDUCATION",
        "EMPLOYMENT HISTORY",
        "ADDITIONAL INFORMATION",
    ]:
        p = doc.add_paragraph(title)
        try:
            p.style = doc.styles["Heading 2"]
        except KeyError:
            pass
    doc.save(path)

def _sample_fields():
    # minimal but realistic data to exercise writers
    return {
        "summary": "Seasoned, results-driven leader with international experience across strategy, ops, and finance.",
        "skills": ["Leadership", "Financial Modelling"],
        "professional_development": [
            "LeanSigma Champion Training, TBM",
            "EVA Training, Stern Stewart",
        ],
        "education": [
            {"institution": "Uni Somewhere", "award": "Distinction", "qualification": "MBA (Executive)", "year": "2020"},
        ],
        "experience": [
            {
                "company": "Franchise Brands plc, London",
                "title": "Head of Ops",
                "company_info": "Multi-brand franchisor.",
                "bullets": ["Led X", "Saved £3m"],
                "dates": "2019–Present",
            }
        ],
        "additional_information": ["Full UK Driving Licence"],
    }

def _has_nonempty_body_after(doc: Document, heading_text: str) -> bool:
    """Find heading and check at least one following paragraph has text."""
    idx = None
    for i, p in enumerate(doc.paragraphs):
        if (p.text or "").strip().upper() == heading_text.upper():
            idx = i
            break
    if idx is None:
        return False
    # scan forward until next heading or end
    for q in doc.paragraphs[idx + 1:]:
        style_name = (getattr(q.style, "name", "") or "").upper()
        if style_name in ("HEADING 1", "HEADING 2", "HEADING 3"):
            break
        if (q.text or "").strip():
            return True
    return False

def test_pipeline_e2e_builds_doc_with_correct_section_order(tmp_path):
    template = tmp_path / "template.docx"
    _make_minimal_template(str(template))

    data_path = tmp_path / "fields.json"
    with open(data_path, "w", encoding="utf-8") as f:
        json.dump(_sample_fields(), f)

    out_path = tmp_path / "out.docx"

    reformat_cv_cv1_to_template1(
        input_docx=str(template),
        template_docx=str(template),
        out_path=str(out_path),
        data_json=str(data_path),
        meta_profile="template1",
        no_legacy=True,
        section_profile_name="template1_sections.json",
    )

    doc = Document(str(out_path))
    joined = "\n".join(p.text for p in doc.paragraphs)

    # 1) All headings present
    order = [
        "PERSONAL PROFILE",
        "KEY SKILLS",
        "PROFESSIONAL DEVELOPMENT",
        "EDUCATION",
        "EMPLOYMENT HISTORY",
        "ADDITIONAL INFORMATION",
    ]
    idx = {k: joined.find(k) for k in order}
    assert all(v >= 0 for v in idx.values())

    # 2) In canonical order
    assert idx["PERSONAL PROFILE"] < idx["KEY SKILLS"] < idx["PROFESSIONAL DEVELOPMENT"] < idx["EDUCATION"] < idx["EMPLOYMENT HISTORY"] < idx["ADDITIONAL INFORMATION"]

    # 3) Each section has at least one non-empty paragraph written by its writer
    for h in order:
        assert _has_nonempty_body_after(doc, h), f"No body text under heading: {h}"
