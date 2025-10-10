# tests/test_experience_writer.py
import os
from docx import Document
from cv_reformatter.sections import experience as exp

def _sample_data():
    return {
        "employment_history": [
            {
                "title": "Head of Operations",
                "company": "Franchise Brands plc",
                "location": "London",
                "start": "Jan 2021",
                "end": "Present",
                "company_blurb": "Multi-brand franchisor in B2B services.",
                "bullets": [
                    "Led cross-functional transformation programme.",
                    "Delivered £3.2m cost savings via lean initiatives."
                ],
            },
            {
                "title": "Strategy Director",
                "company": "Pirtek Europe",
                "location": "Birmingham",
                "start": "2017",
                "end": "2020",
                "bullets": [
                    "Devised market entry strategy across 3 regions.",
                ],
            },
        ]
    }

def test_experience_writer_renders_heading_and_blocks(tmp_path):
    doc = Document()
    data = _sample_data()

    exp.write_section(doc, "EMPLOYMENT HISTORY", "", data)

    texts = [p.text for p in doc.paragraphs]
    # Heading present
    assert any(t.strip() == "EMPLOYMENT HISTORY" for t in texts)

    # First role lines present
    assert any("Jan 2021" in t and "Present" in t and "Franchise Brands plc" in t for t in texts)
    assert any("Head of Operations" == t.strip() for t in texts)

    # Company blurb italic
    italics = []
    for p in doc.paragraphs:
        if "Multi-brand franchisor in B2B services." in p.text:
            italics = [run.italic for run in p.runs]
            break
    assert italics and any(italics), "Company blurb should be italic"

    # Bullets styled (List Bullet) - best effort check
    bullet_count = 0
    for p in doc.paragraphs:
        if p.text.strip() in {
            "Led cross-functional transformation programme.",
            "Delivered £3.2m cost savings via lean initiatives.",
            "Devised market entry strategy across 3 regions.",
        }:
            bullet_count += 1
            # style name may vary slightly across Word locales; allow fallback
            assert ("List Bullet" in str(p.style.name)) or p.style.name.lower().startswith("list")
    assert bullet_count == 3
