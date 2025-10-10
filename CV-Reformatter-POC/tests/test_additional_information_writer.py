from docx import Document
from cv_reformatter.sections import extras as ex

def test_additional_information_present_when_data_exists():
    doc = Document()
    data = {"additional_information": ["Full UK Driving Licence", "Willing to travel"]}
    ex.write_section(doc, "ADDITIONAL INFORMATION", "", data)

    lines = [p.text.strip() for p in doc.paragraphs]
    assert "ADDITIONAL INFORMATION" in lines
    assert "Full UK Driving Licence" in lines
    assert "Willing to travel" in lines

    # bullets styled (best-effort)
    for p in doc.paragraphs:
        if p.text.strip() in {"Full UK Driving Licence", "Willing to travel"}:
            assert ("List" in p.style.name) or p.style.name.lower().startswith("list")
