# tests/test_education_writer.py
from docx import Document
from cv_reformatter.sections import education as edu

def _sample_data():
    return {
        "education": [
            {
                "year": "2021",
                "institution": "University of Somewhere",
                "title": "MBA (Executive)",
                "award": "Distinction",
                "bullets": ["Corporate Finance I & II", "International Finance"]
            },
            {
                "year": "2018",
                "institution": "Tech University",
                "title": "BEng (Hons) Mechanical Engineering",
                "award": "First Class",
                "bullets": ["Capstone project on lean design"]
            }
        ]
    }

def test_education_writer_renders_heading_and_entries_in_order():
    doc = Document()
    data = _sample_data()

    edu.write_section(doc, "EDUCATION", "", data)
    lines = [p.text for p in doc.paragraphs]

    # Heading
    assert any(l.strip() == "EDUCATION" for l in lines)

    # First entry (2021) should come before 2018
    joined = "\n".join(lines)
    idx_2021 = joined.find("2021")
    idx_2018 = joined.find("2018")
    assert idx_2021 != -1 and idx_2018 != -1 and idx_2021 < idx_2018

    # Check tab-separated line structure for 2021 entry
    assert any(("2021" in l and "\t" in l and "University of Somewhere" in l) for l in lines)

    # Title + award line present
    assert any("MBA (Executive)" in l and "Distinction" in l for l in lines)

    # Bullets styled (best-effort: style name contains 'List')
    bullet_texts = {
        "Corporate Finance I & II",
        "International Finance",
        "BEng (Hons) Mechanical Engineering",
        "First Class",
        "Capstone project on lean design"
    }
    seen = 0
    for p in doc.paragraphs:
        if p.text.strip() in bullet_texts:
            seen += 1
            assert ("List" in p.style.name) or (p.style.name.lower().startswith("list"))
    assert seen >= 1  # at least one bullet should be styled as a list
