import hashlib
from docx import Document
from docx.shared import Pt
from cv_reformatter.meta import apply_meta_with_profile

def _mk_profile(*, lines=2, name="Calibri", size=12, bold=True, all_caps=True):
    return {"title_block": {
        "apply": True, "lines": lines, "name": name, "size": size,
        "bold": bold, "all_caps": all_caps
    }}

def _get_run_attrs(p):
    out = []
    for r in p.runs:
        fn = r.font.name
        sz = r.font.size.pt if r.font.size else None
        out.append((fn, sz, bool(r.font.bold), bool(r.font.all_caps)))
    return out

def _fingerprint(doc: Document) -> str:
    import io, zipfile
    bio = io.BytesIO()
    doc.save(bio)
    z = zipfile.ZipFile(io.BytesIO(bio.getvalue()))
    parts = []
    for info in sorted(z.infolist(), key=lambda i: i.filename):
        parts.append(info.filename.encode("utf-8"))
        parts.append(z.read(info.filename))
    return hashlib.md5(b"".join(parts)).hexdigest()

def test_title_block_applies_first_two_paragraphs_only_and_idempotent():
    doc = Document()
    # 3 paragraphs to check the 3rd is left untouched
    p1 = doc.add_paragraph("FIRST LINE")
    p2 = doc.add_paragraph("SECOND LINE")
    p3 = doc.add_paragraph("THIRD LINE (should not be styled as title)")

    prof = _mk_profile(lines=2, name="Calibri", size=12, bold=True, all_caps=True)
    apply_meta_with_profile(doc, prof)

    # First two paragraphs styled as a single run with Calibri 12 bold all_caps
    attrs1 = _get_run_attrs(p1)
    attrs2 = _get_run_attrs(p2)
    assert len(attrs1) == 1 and len(attrs2) == 1
    assert attrs1[0][0] == "Calibri" and attrs2[0][0] == "Calibri"
    assert attrs1[0][1] == 12 and attrs2[0][1] == 12
    assert attrs1[0][2] is True and attrs2[0][2] is True
    assert attrs1[0][3] is True and attrs2[0][3] is True

    # Third paragraph unchanged (not forced into a single styled run)
    # It will still be one run from python-docx, but not styled by our code.
    # We assert its size is NOT 12 as a quick check (could be None).
    a3 = _get_run_attrs(p3)
    assert len(a3) == 1
    assert a3[0][1] != 12  # not styled like the title block

    # Idempotent on re-apply
    before = _fingerprint(doc)
    apply_meta_with_profile(doc, prof)
    after = _fingerprint(doc)
    assert before == after
