from docx import Document
from cv_reformatter.sections import skills as sk

def _data(skills):
    return {"skills": skills}

def test_skills_renders_heading_and_bullets():
    doc = Document()
    data = _data(["Leadership", "Financial Modelling", "Lean / Kaizen"])
    sk.write_section(doc, "KEY SKILLS", "", data)

    lines = [p.text.strip() for p in doc.paragraphs]
    assert "KEY SKILLS" in lines
    # bullets present
    for needed in ["Leadership", "Financial Modelling", "Lean / Kaizen"]:
        assert needed in lines

    # best-effort: style contains 'List' for bullet lines
    for p in doc.paragraphs:
        if p.text.strip() in {"Leadership", "Financial Modelling", "Lean / Kaizen"}:
            assert ("List" in p.style.name) or p.style.name.lower().startswith("list")
