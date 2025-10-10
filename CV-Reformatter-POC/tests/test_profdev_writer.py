from docx import Document
from cv_reformatter.sections import extras as ex

def _data(pd_items):
    # mimic remap from 'Other Headings' too
    return {"professional_development": pd_items, "other_headings": []}

def test_professional_development_heading_and_bullets():
    doc = Document()
    data = _data([
        "LeanSigma Champion Training, TBM",
        "EVA Training, Stern Stewart",
        "MBA Electives: Advanced Corporate Finance I & II"
    ])

    ex.write_section(doc, "PROFESSIONAL DEVELOPMENT", "", data)
    lines = [p.text.strip() for p in doc.paragraphs]
    assert "PROFESSIONAL DEVELOPMENT" in lines

    for needed in [
        "LeanSigma Champion Training, TBM",
        "EVA Training, Stern Stewart",
        "MBA Electives: Advanced Corporate Finance I & II"
    ]:
        assert needed in lines

    for p in doc.paragraphs:
        if p.text.strip() in {
            "LeanSigma Champion Training, TBM",
            "EVA Training, Stern Stewart",
            "MBA Electives: Advanced Corporate Finance I & II"
        }:
            assert ("List" in p.style.name) or p.style.name.lower().startswith("list")
