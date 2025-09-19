import argparse
from pathlib import Path
from docx import Document

PLACEHOLDERS = [
    "NAME", "EMAIL", "PHONE", "URL",
    "SUMMARY", "EXPERIENCE", "EDUCATION", "SKILLS"
]

def _docx_has_placeholders(doc):
    tokens = {"{{" + k + "}}" for k in PLACEHOLDERS}
    for p in doc.paragraphs:
        t = p.text or ""
        if any(tok in t for tok in tokens):
            return True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text or ""
                    if any(tok in t for tok in tokens):
                        return True
    return False

def _append_placeholder_page(doc):
    doc.add_page_break()
    doc.add_heading("Template Tokens (POC)", level=1)
    def add_kv(label, key):
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run("{{" + key + "}}")
    add_kv("Name", "NAME"); add_kv("Email", "EMAIL"); add_kv("Phone", "PHONE"); add_kv("URL", "URL")
    doc.add_paragraph(" "); doc.add_heading("Summary", level=2); doc.add_paragraph("{{SUMMARY}}")
    doc.add_paragraph(" "); doc.add_heading("Experience", level=2); doc.add_paragraph("{{EXPERIENCE}}")
    doc.add_paragraph(" "); doc.add_heading("Education", level=2); doc.add_paragraph("{{EDUCATION}}")
    doc.add_paragraph(" "); doc.add_heading("Skills", level=2); doc.add_paragraph("{{SKILLS}}")

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", required=True)
    ap.add_argument("--out", required=True)
    args = ap.parse_args()

    tmpl = Path(args.template)
    outp = Path(args.out); outp.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(tmpl))

    if _docx_has_placeholders(doc):
        print("[ensure_placeholders] Template already has placeholders. Saving copy unchanged.")
    else:
        print("[ensure_placeholders] No placeholders found. Appending a token page...")
        _append_placeholder_page(doc)

    doc.save(str(outp))
    print(f"[ensure_placeholders] Wrote: {outp}")

if __name__ == "__main__":
    main()
