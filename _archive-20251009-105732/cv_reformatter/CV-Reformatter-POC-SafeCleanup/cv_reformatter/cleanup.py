
from typing import Dict, Any
from docx.text.paragraph import Paragraph

def _norm(s: str) -> str:
    return " ".join((s or "").replace("\xa0", " ").strip().split()).upper()

def _remove_paragraph(p: Paragraph) -> None:
    p._element.getparent().remove(p._element)

def _row_text(row) -> str:
    try:
        return _norm("\n".join(c.text for c in row.cells))
    except Exception:
        return ""

def _looks_like_qual_guidance(text: str) -> bool:
    t = text.upper()
    return (
        "NAME OF ESTABLISHMENT" in t or
        "TITLE OF QUALIFICATION" in t or
        "DATE" == t or
        "LIST MOST RECENT FIRST." in t
    )

def prune_placeholders(doc, data: Dict[str, Any]) -> None:
    # Only act if there is no qualifications data
    if data.get("qualifications"):
        return

    # Prefer table-row removal (most templates place this in a row)
    for tbl in doc.tables:
        rows = list(tbl.rows)
        for i, row in enumerate(rows):
            if _row_text(row) == "QUALIFICATIONS":
                try:
                    tbl._tbl.remove(row._tr)  # remove heading row
                    # Remove the next row if it's clearly guidance
                    if i < len(tbl.rows):
                        next_rt = _row_text(tbl.rows[i])
                        if _looks_like_qual_guidance(next_rt):
                            tbl._tbl.remove(tbl.rows[i]._tr)
                except Exception:
                    pass
                return  # Done after first match

    # Fallback: non-table paragraph heading
    for idx, p in enumerate(list(doc.paragraphs)):
        if _norm(p.text) == "QUALIFICATIONS":
            _remove_paragraph(p)
            # Optional: remove one following guidance paragraph if it looks like guidance
            if idx < len(doc.paragraphs):
                q = doc.paragraphs[idx]
                if _looks_like_qual_guidance(_norm(q.text)):
                    _remove_paragraph(q)
            return
