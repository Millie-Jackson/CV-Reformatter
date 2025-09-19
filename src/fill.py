"""
fill.py
-------
Template filler for DOCX using placeholder tokens like:
  {{NAME}}, {{EMAIL}}, {{PHONE}}, {{URL}},
  {{SUMMARY}}, {{EXPERIENCE}}, {{EDUCATION}}, {{SKILLS}}

This POC replaces placeholders across all paragraphs and table cells.
If a placeholder isn't present in the template, it's simply skipped.
"""

from typing import Dict, Any
from docx import Document


PLACEHOLDERS = [
    "NAME", "EMAIL", "PHONE", "URL",
    "SUMMARY", "EXPERIENCE", "EDUCATION", "SKILLS"
]


def _collapse_ws(text: str) -> str:
    import re
    return re.sub(r"\s+", " ", (text or "").strip())


def build_mapping(fields: Dict[str, Any]) -> Dict[str, str]:
    """Convert extracted fields to string mapping for placeholders."""
    def val(x): return _collapse_ws(x) if isinstance(x, str) else x
    mapping = {
        "NAME": val(fields.get("name")) or "—",
        "EMAIL": val(fields.get("email")) or "—",
        "PHONE": val(fields.get("phone")) or "—",
        "URL": val(fields.get("url")) or "—",
        "SUMMARY": val(fields.get("summary")) or "—",
        "EXPERIENCE": val(fields.get("experience_raw")) or "—",
        "EDUCATION": val(fields.get("education_raw")) or "—",
        "SKILLS": ", ".join(fields.get("skills") or []) or "—",
    }
    return mapping


def _replace_in_text(original: str, mapping: Dict[str, str]) -> str:
    text = original
    for key, val in mapping.items():
        token = "{{" + key + "}}"
        text = text.replace(token, val)
    return text


def _replace_in_paragraph(paragraph, mapping: Dict[str, str]) -> bool:
    original = paragraph.text
    new = _replace_in_text(original, mapping)
    if new == original:
        return False
    # wipe existing runs and replace with a single run to keep it simple for POC
    for run in paragraph.runs:
        run.text = ""
    paragraph.add_run(new)
    return True


def _replace_in_table(table, mapping: Dict[str, str]) -> int:
    count = 0
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if _replace_in_paragraph(p, mapping):
                    count += 1
    return count


def fill_template(template_path: str, output_path: str, mapping: Dict[str, str]) -> Dict[str, int]:
    """
    Replace placeholders throughout the document.
    Returns a dict with counts of replacements made.
    """
    doc = Document(template_path)
    replaced = 0

    # paragraphs
    for p in doc.paragraphs:
        if _replace_in_paragraph(p, mapping):
            replaced += 1

    # tables
    for t in doc.tables:
        replaced += _replace_in_table(t, mapping)

    doc.save(output_path)
    return {"replacements": replaced}
