# autofill.py (stronger cleanup of 'Xxxxx…' placeholders)
import re
from typing import Dict, Optional, List, Tuple

from docx import Document
from docx.text.paragraph import Paragraph

from utilities import (
    letter_space_two,
    ensure_font_calibri_10,
    apply_style_if_exists,
    new_paragraph_after,
    add_blank_lines_before,
    normalise_punctuation,
)

LABELS = {
    "summary":    re.compile(r"(?i)\b(profile|personal\s+profile|summary|objective|about\s+me)\b"),
    "experience": re.compile(r"(?i)\b(experience|employment|employment\s+history|work\s+history|career\s+history)\b"),
    "education":  re.compile(r"(?i)\b(education|education\s*&\s*qualifications|education\s+and\s+qualifications|qualifications|academic)\b"),
    "skills":     re.compile(r"(?i)\b(skills|key\s+skills|technical\s+skills|core\s+skills|skills\s*(?:&|and)\s*competencies|competencies)\b"),
}

FILLER_PATTERNS = [
    re.compile(r"^other\s+headings?\b", re.I),
    re.compile(r"^x{5,}$", re.I),
    re.compile(r"^list most recent first\.?$", re.I),
    re.compile(r"^start\s*date", re.I),
    re.compile(r"^date\s*[-–—]?\s*", re.I),
    re.compile(r"^job\s*title$", re.I),
    re.compile(r"^company", re.I),
    re.compile(r"^<insert .*?>", re.I),
    re.compile(r"^[–—-]\s*$"),
    re.compile(r"^\s+$"),
]

CV_BODY = "CV Body"
CV_BULLET = "CV Bullet"

def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip()).lower()

def _replace_paragraph_text(paragraph: Paragraph, text: str, bold: bool = False) -> None:
    for r in paragraph.runs:
        r.text = ""
        r.bold = False
    run = paragraph.add_run(text or "")
    run.bold = bool(bold)
    ensure_font_calibri_10(paragraph)

def _find_heading(doc, section_key: str) -> Optional[Paragraph]:
    for p in doc.paragraphs:
        t = _norm(p.text)
        style = getattr(getattr(p, "style", None), "name", "") or ""
        if LABELS[section_key].search(t) or ("heading" in style.lower() and LABELS[section_key].search(t)):
            return p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = _norm(p.text)
                    style = getattr(getattr(p, "style", None), "name", "") or ""
                    if LABELS[section_key].search(t) or ("heading" in style.lower() and LABELS[section_key].search(t)):
                        return p
    return None

def _find_title_paragraph(doc) -> Optional[Paragraph]:
    for p in doc.paragraphs[:12]:
        style = getattr(getattr(p, "style", None), "name", "") or ""
        if "title" in style.lower():
            return p
    for p in doc.paragraphs[:20]:
        if re.search(r"(?i)\bcurriculum\s+vitae\b", p.text):
            return p
    return doc.paragraphs[0] if doc.paragraphs else None

def _find_location_placeholders(doc) -> List[Paragraph]:
    nodes = []
    rx = re.compile(r"(?i)\bcandidate\s+location\b")
    for p in doc.paragraphs[:100]:
        if rx.search(p.text):
            nodes.append(p)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if rx.search(p.text):
                        nodes.append(p)
    return nodes

DATE_RANGE_RE = re.compile(r"(?i)\b(?:\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}|\bjan|\bfeb|\bmar|\bapr|\bmay|\bjun|\bjul|\baug|\bsep|\boct|\bnov|\bdec|\d{4})[^\n]{0,30}?\b(?:present|current|\d{4})\b")

def _split_experience_entries(text: str) -> List[str]:
    blocks = re.split(r"\n\s*\n", (text or '').strip())
    merged, cur = [], []
    for b in blocks:
        if DATE_RANGE_RE.search(b) and cur:
            merged.append("\n".join(cur))
            cur = [b]
        else:
            cur.append(b)
    if cur:
        merged.append("\n".join(cur))
    return [m.strip() for m in merged if m.strip()]

def _format_company_header(p: Paragraph, company: str, title: str, dates: str, info: str = "") -> None:
    add_blank_lines_before(p, 1)
    for r in p.runs:
        r.text = ""
        r.bold = False
    p.add_run((dates or "").strip()).bold = True
    p.add_run("  ")
    p.add_run((company or "").strip().upper()).bold = True
    p.add_run(" — ")
    p.add_run((title or "").strip().title()).bold = True
    if info:
        p.add_run(" — " + info.strip())
    apply_style_if_exists(p, CV_BODY) or apply_style_if_exists(p, "Normal")
    ensure_font_calibri_10(p)

def _insert_paragraph_after(heading_p: Paragraph, text: str, bullet: bool = False) -> Paragraph:
    p = new_paragraph_after(heading_p)
    cleaned = re.sub(r"^\s*[•\-–—]\s*", "", text or "")
    p.add_run(normalise_punctuation(cleaned or " "))
    if bullet:
        apply_style_if_exists(p, CV_BULLET) or apply_style_if_exists(p, "List Bullet")
    else:
        apply_style_if_exists(p, CV_BODY) or apply_style_if_exists(p, "Normal")
    ensure_font_calibri_10(p)
    return p

# --- NEW: robust placeholder detector (handles 'Xxxxxxxxxxx.' bullets etc.) ---
def _is_placeholder_text(s: str) -> bool:
    if not s:
        return True
    t = (s or "").strip()
    # Strip common punctuation/spaces
    core = re.sub(r"[\s\.,;:!\-–—_]+", "", t)
    if not core:
        return True
    # If 80%+ of remaining chars are 'x' (any case), treat as placeholder
    x_count = sum(1 for ch in core if ch.lower() == "x")
    if x_count / len(core) >= 0.8 and len(core) >= 5:
        return True
    # Obvious keywords
    if re.search(r"(?i)start\s*date|job\s*title|most recent", t):
        return True
    return False

def _clear_filler_after(heading_p: Paragraph, max_to_clear: int = 80) -> None:
    from docx.text.paragraph import Paragraph as P
    el = heading_p._p.getnext()
    cleared = 0
    while el is not None and cleared < max_to_clear:
        tag = el.tag.lower()
        if tag.endswith("tbl"):
            tbl = el
            text_fragments = []
            for row in tbl.xpath(".//w:tr", namespaces=tbl.nsmap):
                for cell in row.xpath(".//w:tc", namespaces=tbl.nsmap):
                    for p in cell.xpath(".//w:p", namespaces=tbl.nsmap):
                        texts = [t.text for t in p.xpath(".//w:t", namespaces=p.nsmap) if t.text]
                        if texts:
                            text_fragments.append("".join(texts))
            table_text = "\n".join(text_fragments)
            if _is_placeholder_text(table_text):
                parent = el.getparent(); parent.remove(el); cleared += 1
                el = heading_p._p.getnext(); continue
            else:
                break
        elif tag.endswith("p"):
            from docx.text.paragraph import Paragraph as P
            para = P(el, heading_p._parent)
            txt = (para.text or "").strip()
            style_name = getattr(getattr(para, "style", None), "name", "") or ""
            if "heading" in style_name.lower() or "title" in style_name.lower():
                break
            if _is_placeholder_text(txt) or any(rx.search(txt) for rx in FILLER_PATTERNS):
                parent = para._element.getparent(); parent.remove(para._element)
                cleared += 1; el = heading_p._p.getnext(); continue
            break
        else:
            break

def _purge_placeholders_in_section(head_p: Paragraph) -> None:
    """Remove remaining placeholder paragraphs/tables until the next heading."""
    from docx.text.paragraph import Paragraph as P
    el = head_p._p.getnext()
    while el is not None:
        tag = el.tag.lower()
        if tag.endswith("p"):
            para = P(el, head_p._parent)
            style_name = getattr(getattr(para, "style", None), "name", "") or ""
            if "heading" in style_name.lower() or "title" in style_name.lower():
                break
            txt = (para.text or "").strip()
            if _is_placeholder_text(txt) or any(rx.search(txt) for rx in FILLER_PATTERNS):
                parent = para._element.getparent(); parent.remove(para._element)
                el = head_p._p.getnext(); continue
        elif tag.endswith("tbl"):
            tbl = el
            text_fragments = []
            for row in tbl.xpath(".//w:tr", namespaces=tbl.nsmap):
                for cell in row.xpath(".//w:tc", namespaces=tbl.nsmap):
                    for p in cell.xpath(".//w:p", namespaces=tbl.nsmap):
                        texts = [t.text for t in p.xpath(".//w:t", namespaces=p.nsmap) if t.text]
                        if texts: text_fragments.append("".join(texts))
            if _is_placeholder_text("\n".join(text_fragments)):
                parent = el.getparent(); parent.remove(el)
                el = head_p._p.getnext(); continue
        el = el.getnext()


def _remove_global_template_labels(doc: Document) -> None:
    """Remove generic template notes like 'OTHER HEADINGS' anywhere in the doc (paragraphs or table cells)."""
    import docx
    rx = re.compile(r"(?i)^\s*other\s+headings?\s*$")
    # Paragraphs
    for p in list(doc.paragraphs):
        if rx.match((p.text or "").strip()):
            parent = p._element.getparent()
            parent.remove(p._element)
    # Tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in list(cell.paragraphs):
                    if rx.match((p.text or "").strip()):
                        par_el = p._element
                        par_el.getparent().remove(par_el)
def autofill_by_labels(template_path: str, output_path: str, mapping: Dict[str, str], meta: Optional[Dict[str, str]] = None) -> Dict[str, int]:
    doc = Document(template_path)
    changes = 0

    name = (mapping.get("NAME") or "").strip()
    title_p = _find_title_paragraph(doc)
    header_text = f"CURRICULUM VITAE FOR {name}" if name else "CURRICULUM VITAE"
    if title_p is not None:
        _replace_paragraph_text(title_p, letter_space_two(header_text), bold=True); changes += 1

    location = (mapping.get("LOCATION") or "CANDIDATE LOCATION").strip()
    loc_nodes = _find_location_placeholders(doc)
    if loc_nodes:
        for node in loc_nodes:
            _replace_paragraph_text(node, letter_space_two(location), bold=True); changes += 1
    else:
        if title_p is not None:
            after = new_paragraph_after(title_p)
            _replace_paragraph_text(after, letter_space_two(location), bold=True); changes += 1

    if mapping.get("SUMMARY"):
        p_sum = _find_heading(doc, "summary") or doc.add_heading("Personal Profile", level=2)
        _clear_filler_after(p_sum)
        for line in (mapping["SUMMARY"] or "").splitlines():
            p_sum = _insert_paragraph_after(p_sum, line, bullet=False)
        _purge_placeholders_in_section(p_sum)
        changes += 1

    if mapping.get("EXPERIENCE"):
        p_head = _find_heading(doc, "experience") or doc.add_heading("Employment History", level=2)
        _clear_filler_after(p_head)
        _purge_placeholders_in_section(p_head)
        entries = _split_experience_entries(mapping["EXPERIENCE"])
        anchor = p_head
        for entry in entries:
            lines = [l for l in entry.splitlines() if l.strip()]
            if not lines:
                continue
            header = lines[0]
            m = re.search(DATE_RANGE_RE, header)
            dates = m.group(0) if m else ""
            core = header
            if dates:
                core = core.replace(dates, "").strip(" -—–,\t:")
            company = core.split("—")[0].split("-")[0].strip()
            title = core.replace(company, "", 1).strip(" -—–,\t")

            info_lines = []
            i = 1
            while i < len(lines):
                t = lines[i].strip()
                if not t:
                    i += 1; continue
                if re.match(r"(?i)^responsibilit(y|ies)\s*:", t) or re.match(r"^\s*[•\-–—]\s*", t):
                    break
                if len(t) <= 80:
                    info_lines.append(t); i += 1; continue
                break
            if i < len(lines) and re.match(r"(?i)^responsibilit(y|ies)\s*:", lines[i].strip()):
                i += 1
            body_lines = [l for l in lines[i:] if l.strip()]
            info = ", ".join(info_lines)

            header_p = new_paragraph_after(anchor)
            _format_company_header(header_p, company, title, dates, info=info); changes += 1

            for bl in body_lines:
                if re.match(r"(?i)^responsibilit(y|ies)\s*:", bl.strip()):
                    continue
                header_p = _insert_paragraph_after(header_p, bl, bullet=True)
            anchor = header_p
        _purge_placeholders_in_section(p_head)

    if mapping.get("SKILLS"):
        p_sk = _find_heading(doc, "skills") or doc.add_heading("Key Skills", level=2)
        _clear_filler_after(p_sk)
        for skill in (mapping["SKILLS"] or "").splitlines():
            if not skill.strip():
                continue
            p_sk = _insert_paragraph_after(p_sk, skill, bullet=True)
        _purge_placeholders_in_section(p_sk)
        changes += 1

    if mapping.get("EDUCATION"):
        p_edu = _find_heading(doc, "education") or doc.add_heading("Education", level=2)
        _clear_filler_after(p_edu)
        for line in (mapping["EDUCATION"] or "").splitlines():
            p_edu = _insert_paragraph_after(p_edu, line, bullet=False)
        _purge_placeholders_in_section(p_edu)
        changes += 1

    _remove_global_template_labels(doc)
    doc.save(output_path)
    return {"changes": changes}
