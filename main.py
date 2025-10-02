#!/usr/bin/env python3
import argparse, json, os, re
from typing import List, Dict, Optional, Union
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ============== IO ==============
def load_json(path: str) -> Dict:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)

def find_template(path_arg: Optional[str]) -> str:
    candidates = [path_arg] if path_arg else []
    candidates += [
        "Templates & Briefs/Template 1.docx",
        "data/templates/Template1.docx",
        "Templates & Briefs/Template 2.docx",
    ]
    for c in candidates:
        if c and os.path.isfile(c):
            return c
    raise FileNotFoundError("Template .docx not found. Pass with -t/--template.")

# ============== Paragraph utils ==============
def delete_paragraph(p):
    p._element.getparent().remove(p._element)

def insert_paragraph_after(paragraph, text: str = ""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph as Para
    p = Para(new_p, paragraph._parent)
    if text:
        p.add_run(text)
    return p

def insert_paragraph_before(paragraph, text: str = ""):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    from docx.text.paragraph import Paragraph as Para
    p = Para(new_p, paragraph._parent)
    if text:
        p.add_run(text)
    return p

def set_run(run, *, name="Calibri", size_pt=10, bold=False, italic=False, all_caps=False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.all_caps = all_caps

def format_paragraph(p, *, font_name="Calibri", font_size_pt=10, bold=False, italic=False, all_caps=False, space_after_pt=0):
    if not p.runs:
        r = p.add_run("")
        set_run(r, name=font_name, size_pt=font_size_pt, bold=bold, italic=italic, all_caps=all_caps)
    else:
        for r in p.runs:
            set_run(r, name=font_name, size_pt=font_size_pt, bold=bold, italic=italic, all_caps=all_caps)
    p.paragraph_format.space_after = Pt(space_after_pt)

def find_paragraph(doc: Document, predicate) -> Optional[int]:
    for i, p in enumerate(doc.paragraphs):
        if predicate(p.text):
            return i
    return None

def find_heading_index(doc: Document, *candidates: str) -> Optional[int]:
    cand_low = [c.lower() for c in candidates]
    return find_paragraph(doc, lambda t: t.strip().lower() in cand_low)

def startswith_line(doc: Document, text_prefix: str) -> Optional[int]:
    pref = text_prefix.lower()
    return find_paragraph(doc, lambda t: t.strip().lower().startswith(pref))

def is_heading_text(text: str) -> bool:
    t = (text or "").strip()
    return bool(t) and t.isupper() and len(t) > 3

# ============== Heading styles & helpers ==============
def _apply_heading_borders(p):
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    def ensure_edge(edge):
        el = pBdr.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            pBdr.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '1')
        el.set(qn('w:color'), '4F81BD')
    ensure_edge('top'); ensure_edge('bottom')

def apply_heading_style(doc: Document, heading_text: str):
    idx = find_heading_index(doc, heading_text)
    if idx is not None:
        p = doc.paragraphs[idx]
        p.text = heading_text.upper()
        if not p.runs:
            p.add_run("")
        for r in p.runs:
            set_run(r, name="Calibri", size_pt=10, bold=False, italic=True, all_caps=True)
            r.font.color.rgb = RGBColor(79, 129, 189)
        p.paragraph_format.space_after = Pt(0)

def clear_paragraph_borders(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is not None:
        pPr.remove(pBdr)

# ============== Spacing helpers ==============
def set_space_before(p, pts: int):
    try:
        p.paragraph_format.space_before = Pt(pts)
    except Exception:
        pass

def set_space_after(p, pts: int):
    try:
        p.paragraph_format.space_after = Pt(pts)
    except Exception:
        pass

def tighten_to_next_heading(doc: "Document", current_heading: str, next_head_space_before_pt: int = 12):
    """
    Remove blank paragraphs before the next heading; set small space-before on that heading.
    """
    h_idx = find_heading_index(doc, current_heading)
    if h_idx is None:
        return
    # locate next heading
    n_idx = None
    for i in range(h_idx + 1, len(doc.paragraphs)):
        t = (doc.paragraphs[i].text or "").strip()
        if is_heading_text(t):
            n_idx = i
            break
    if n_idx is None:
        return
    # delete any blanks directly before the next heading
    i = n_idx - 1
    while i > h_idx and (doc.paragraphs[i].text or "").strip() == "":
        delete_paragraph(doc.paragraphs[i])
        n_idx -= 1
        i -= 1
    # apply small space-before on the next heading
    try:
        doc.paragraphs[n_idx].paragraph_format.space_before = Pt(next_head_space_before_pt)
    except Exception:
        pass

# ============== Normalisers ==============
def _year_key(dates: str) -> int:
    m = re.findall(r"(19|20)\d{2}", dates or "")
    return int(m[-1]) if m else -1

def _normalise_skills(skills: Union[Dict, List, str, None]) -> List[str]:
    items: List[str] = []
    if skills is None:
        return items

    def normalise_chunk(s: str) -> str:
        if not s:
            return ""
        s = re.sub(r"-\s*[\r\n]+\s*", "", s)
        s = re.sub(r"[\r\n]+", " ", s)
        s = re.sub(r"\s+", " ", s)
        return s.strip(" •-\t.")

    def extend_from_text(s: str):
        if not s:
            return
        parts = re.split(r"\n{2,}|;", s)
        for part in parts:
            clean = normalise_chunk(part)
            if clean:
                items.append(clean)

    if isinstance(skills, list):
        for x in skills:
            extend_from_text(x if isinstance(x, str) else str(x))
    elif isinstance(skills, dict):
        for v in skills.values():
            if isinstance(v, list):
                for x in v:
                    extend_from_text(x if isinstance(x, str) else str(x))
            else:
                extend_from_text(v if isinstance(v, str) else str(v))
    elif isinstance(skills, str):
        extend_from_text(skills)
    else:
        norm = normalise_chunk(str(skills))
        if norm:
            items.append(norm)

    seen = set(); out = []
    for s in items:
        k = s.lower()
        if s and k not in seen:
            seen.add(k); out.append(s)
    return out

def _coerce_edu_item(x):
    """
    Accept dicts or strings for education/qualifications.
    Returns dict with keys: dates, institution, degree, result.
    Best-effort parse if given a string.
    """
    if isinstance(x, dict):
        return {
            "dates": (x.get("dates") or "").strip(),
            "institution": (x.get("institution") or x.get("establishment") or "").strip(),
            "degree": (x.get("degree") or x.get("title") or "").strip(),
            "result": (x.get("result") or "").strip(),
        }
    s = str(x or "").strip()
    if not s:
        return {"dates":"", "institution":"", "degree":"", "result":""}
    m = re.search(r'(19|20)\d{2}', s)
    dates = m.group(0) if m else ""
    if dates:
        s_wo_year = re.sub(r'(19|20)\d{2}', '', s, count=1).strip(",;- ")
    else:
        s_wo_year = s
    parts = [p.strip() for p in re.split(r'[|,–-]+', s_wo_year) if p.strip()]
    institution = parts[0] if parts else ""
    degree = parts[1] if len(parts) > 1 else ""
    result = parts[2] if len(parts) > 2 else ""
    return {"dates": dates, "institution": institution, "degree": degree, "result": result}

# ============== Heading insert/ensure ==============
def _find_first_heading_index(doc: Document, aliases: list) -> Optional[int]:
    for a in aliases:
        idx = find_heading_index(doc, a)
        if idx is not None:
            return idx
    return None

def _ensure_heading(doc: Document, target: str, aliases: list, before_aliases: list) -> int:
    idx = _find_first_heading_index(doc, [target] + aliases)
    if idx is not None:
        para = doc.paragraphs[idx]
        para.text = target
        apply_heading_style(doc, target)
        _apply_heading_borders(para)
        return idx

    before_idx = _find_first_heading_index(doc, before_aliases) if before_aliases else None
    if before_idx is None:
        from docx.text.paragraph import Paragraph as Para
        new_p = OxmlElement("w:p")
        doc._body._body.append(new_p)
        p_obj = Para(new_p, doc._body)
    else:
        ref_p = doc.paragraphs[before_idx]
        new_p = OxmlElement("w:p")
        ref_p._p.addprevious(new_p)
        from docx.text.paragraph import Paragraph as Para
        p_obj = Para(new_p, ref_p._parent)

    p_obj.text = target
    apply_heading_style(doc, target)
    _apply_heading_borders(p_obj)
    return find_heading_index(doc, target)

# ============== Cleanup helpers ==============
def _strip_placeholders_after_heading(doc: Document, heading_text: str, max_lines: int = 25):
    h = find_heading_index(doc, heading_text)
    if h is None: return
    i = h + 1
    while i < len(doc.paragraphs) and i <= h + max_lines:
        txt = doc.paragraphs[i].text.strip()
        low = txt.lower().replace("–", "-")
        if txt and txt.isupper() and len(txt) > 3:
            break
        placeholders = [
            "list most recent first",
            "educational establishment",
            "awards obtained",
            "name of establishment",
            "title of qualification",
            "date",
            "start date - end date",
            "company info italic",
            "job title",
            "company, location or company info",
            "company, location or company info (most recent roles first)",
        ]
        if (not txt) or any(p in low for p in placeholders):
            delete_paragraph(doc.paragraphs[i]); continue
        else:
            break

def _remove_section(doc: Document, heading_text: str):
    idx = find_heading_index(doc, heading_text)
    if idx is None:
        return
    delete_paragraph(doc.paragraphs[idx])
    while idx < len(doc.paragraphs):
        if idx >= len(doc.paragraphs):
            break
        txt = doc.paragraphs[idx].text.strip()
        if is_heading_text(txt):
            break
        delete_paragraph(doc.paragraphs[idx])

def _remove_template_placeholders(doc: "Document"):
    """
    Remove leftover template instructional lines anywhere in the doc.
    """
    kill_patterns = [
        r"^\s*Start\s*Date\s*[–-]\s*End\s*Date\b",
        r"^\s*COMPANY,?\s*Location\s*or\s*Company\s*Info(\s*\(most\s*recent\s*roles\s*first\))?\s*$",
        r"^\s*Job\s*title\s*$",
        r"^\s*Company\s*Info\s*Italic\s*$",
        r"^\s*OTHER\s*HEADINGS\s*$",
    ]
    compiled = [re.compile(pat, re.IGNORECASE) for pat in kill_patterns]

    i = 0
    while i < len(doc.paragraphs):
        txt = (doc.paragraphs[i].text or "").strip()
        if any(p.search(txt) for p in compiled):
            delete_paragraph(doc.paragraphs[i])
            continue
        i += 1

    while len(doc.paragraphs) > 0 and (doc.paragraphs[-1].text or "").strip() == "":
        delete_paragraph(doc.paragraphs[-1])

def _collapse_multiple_blank_paragraphs(doc: "Document"):
    """
    Collapse any run of 2+ blank paragraphs into a single blank.
    Also removes trailing blanks at the very end.
    """
    i = 0
    prev_blank = False
    while i < len(doc.paragraphs):
        txt = (doc.paragraphs[i].text or "").strip()
        if txt == "":
            if prev_blank:
                delete_paragraph(doc.paragraphs[i])
                continue
            prev_blank = True
        else:
            prev_blank = False
        i += 1
    while len(doc.paragraphs) > 0 and (doc.paragraphs[-1].text or "").strip() == "":
        delete_paragraph(doc.paragraphs[-1])

# ============== Top lines ==============
def style_first_two_lines(doc: "Document", name_line_text: str, location_text: str):
    idx = startswith_line(doc, "CURRICULUM VITAE FOR")
    if idx is not None:
        p = doc.paragraphs[idx]; p.text = ""
        parts = ["CURRICULUM", " ", "VITAE", " ", "FOR", " ", name_line_text.upper() if name_line_text else ""]
        for part in parts:
            r = p.add_run(part)
            set_run(r, name="Calibri", size_pt=12, bold=True, italic=False)
    lidx = startswith_line(doc, "CANDIDATE LOCATION")
    if lidx is not None:
        p = doc.paragraphs[lidx]; p.text = ""
        left = "CANDIDATE LOCATION:"
        right = f" {location_text.upper()}" if location_text else ""
        r1 = p.add_run(left); set_run(r1, name="Calibri", size_pt=12, bold=True, italic=False)
        r2 = p.add_run(right); set_run(r2, name="Calibri", size_pt=12, bold=True, italic=False)

# ============== Bullets ==============
def add_clean_bullet(after_par, text: str):
    """
    Insert a bullet paragraph after `after_par` with tightly controlled spacing/indent.
    Uses a literal bullet and hanging indent.
    """
    bullet = "\u25CF "
    p = insert_paragraph_after(after_par, bullet + (text or ""))
    format_paragraph(p, font_name="Calibri", font_size_pt=10, space_after_pt=0)
    fmt = p.paragraph_format
    fmt.left_indent = Pt(18)       # ~0.63cm
    fmt.first_line_indent = Pt(-9) # hanging indent so wraps align
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    # clear tab stops
    pPr = p._p.get_or_add_pPr()
    tabs = pPr.find(qn('w:tabs'))
    if tabs is not None:
        pPr.remove(tabs)
    # normalize runs
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(10)
    return p

# ============== Writers ==============
def write_header(doc: Document, fields: Dict) -> None:
    name = (fields.get("name") or "").strip()
    location = (fields.get("location") or "").strip()
    style_first_two_lines(doc, name, location)

def write_summary(doc: Document, fields: Dict) -> None:
    repl_idx = find_paragraph(doc, lambda t: "<insert candidate’s executive summary>" in t.lower())
    if repl_idx is not None and fields.get("summary"):
        p = doc.paragraphs[repl_idx]
        p.text = (fields.get("summary") or "").strip()
        format_paragraph(p, font_name="Calibri", font_size_pt=10, bold=False, italic=False, all_caps=False, space_after_pt=0)

def write_skills(doc: Document, fields: Dict) -> None:
    h_idx = find_heading_index(doc, "KEY SKILLS")
    if h_idx is None: return
    apply_heading_style(doc, "KEY SKILLS")
    _strip_placeholders_after_heading(doc, "KEY SKILLS")
    after = doc.paragraphs[h_idx]
    # one blank below heading for breathing room
    first_spacer = insert_paragraph_after(after, "")
    clear_paragraph_borders(first_spacer)
    format_paragraph(first_spacer, font_name="Calibri", font_size_pt=10, space_after_pt=0)
    skills_items = _normalise_skills(fields.get("skills"))
    if not skills_items: return
    prev = first_spacer
    for s in skills_items:
        p = add_clean_bullet(prev, s)
        prev = p
    tighten_to_next_heading(doc, "KEY SKILLS", next_head_space_before_pt=12)

def write_education(doc: Document, education: List[Dict]) -> None:
    apply_heading_style(doc, "EDUCATION")
    h_idx = find_heading_index(doc, "EDUCATION")
    if h_idx is not None:
        _apply_heading_borders(doc.paragraphs[h_idx])

    _strip_placeholders_after_heading(doc, "EDUCATION")
    anchor = doc.paragraphs[find_heading_index(doc, "EDUCATION")]
    after = anchor

    items_raw = education or []
    items = [_coerce_edu_item(x) for x in items_raw]
    items = sorted(items, key=lambda x: _year_key(x.get("dates","")), reverse=True)

    for j, e in enumerate(items):
        year = (e.get("dates") or "").strip()
        inst = (e.get("institution") or "").strip()
        deg  = (e.get("degree") or e.get("title") or "").strip()
        res  = (e.get("result") or "").strip()

        p1 = insert_paragraph_after(after, year + "\t\t\t" + inst)
        format_paragraph(p1, font_name="Calibri", font_size_pt=10, space_after_pt=0)
        if j == 0:
            set_space_before(p1, 6)
        last = p1
        if deg:
            p2 = insert_paragraph_after(p1, "\t\t\t" + deg)
            format_paragraph(p2, font_name="Calibri", font_size_pt=10, space_after_pt=0)
            last = p2
        if res:
            p3 = insert_paragraph_after(last, "\t\t\t" + res)
            format_paragraph(p3, font_name="Calibri", font_size_pt=10, space_after_pt=0)
            last = p3

        after = last
        if j != len(items)-1:
            spacer = insert_paragraph_after(after, "")
            clear_paragraph_borders(spacer)
            format_paragraph(spacer, font_name="Calibri", font_size_pt=10, space_after_pt=0)
            after = spacer

    tighten_to_next_heading(doc, "EDUCATION", next_head_space_before_pt=12)

def write_qualifications(doc: Document, quals: List[Dict]) -> None:
    if not quals:
        _remove_section(doc, "QUALIFICATIONS")
        return
    apply_heading_style(doc, "QUALIFICATIONS")
    h_idx = find_heading_index(doc, "QUALIFICATIONS")
    if h_idx is not None:
        _apply_heading_borders(doc.paragraphs[h_idx])

    _strip_placeholders_after_heading(doc, "QUALIFICATIONS")
    anchor = doc.paragraphs[find_heading_index(doc, "QUALIFICATIONS")]
    after = anchor

    items_raw = quals or []
    items = [_coerce_edu_item(x) for x in items_raw]
    items = sorted(items, key=lambda x: _year_key(x.get("dates","")), reverse=True)

    for j, e in enumerate(items):
        year = (e.get("dates") or "").strip()
        inst = (e.get("institution") or "").strip()
        title  = (e.get("degree") or e.get("title") or "").strip()
        res  = (e.get("result") or "").strip()

        p1 = insert_paragraph_after(after, year + "\t\t\t" + inst)
        format_paragraph(p1, font_name="Calibri", font_size_pt=10, space_after_pt=0)
        if j == 0:
            set_space_before(p1, 6)
        last = p1
        if title:
            p2 = insert_paragraph_after(p1, "\t\t\t" + title)
            format_paragraph(p2, font_name="Calibri", font_size_pt=10, space_after_pt=0)
            last = p2
        if res:
            p3 = insert_paragraph_after(last, "\t\t\t" + res)
            format_paragraph(p3, font_name="Calibri", font_size_pt=10, space_after_pt=0)
            last = p3

        after = last
        if j != len(items)-1:
            spacer = insert_paragraph_after(after, "")
            clear_paragraph_borders(spacer)
            format_paragraph(spacer, font_name="Calibri", font_size_pt=10, space_after_pt=0)
            after = spacer

    tighten_to_next_heading(doc, "QUALIFICATIONS", next_head_space_before_pt=12)

def _write_simple_bullet_section(doc: Document, heading: str, items) -> None:
    if not items:
        _remove_section(doc, heading)
        return

    if heading == "PROFESSIONAL DEVELOPMENT":
        aliases = ["PERSONAL DEVELOPMENT", "DEVELOPMENT", "TRAINING"]
        before_aliases = ["PROFESSIONAL AFFILIATIONS", "EMPLOYMENT HISTORY", "EXPERIENCE", "WORK EXPERIENCE", "PROFESSIONAL EXPERIENCE"]
    elif heading == "PROFESSIONAL AFFILIATIONS":
        aliases = ["AFFILIATIONS", "PROFESSIONAL MEMBERSHIPS", "MEMBERSHIPS"]
        before_aliases = ["EMPLOYMENT HISTORY", "EXPERIENCE", "WORK EXPERIENCE", "PROFESSIONAL EXPERIENCE"]
    else:
        aliases = []
        before_aliases = []

    h_idx = _ensure_heading(doc, heading, aliases, before_aliases)
    _strip_placeholders_after_heading(doc, heading)

    anchor = doc.paragraphs[h_idx]
    spacer = insert_paragraph_after(anchor, "")
    clear_paragraph_borders(spacer)
    format_paragraph(spacer, font_name="Calibri", font_size_pt=10, space_after_pt=0)

    after = spacer
    for it in (items if isinstance(items, list) else [items]):
        txt = re.sub(r"-\s*[\r\n]+\s*", "", str(it))
        txt = re.sub(r"[\r\n]+", " ", txt)
        txt = re.sub(r"\s+", " ", txt).strip(" •-\t.")
        p = add_clean_bullet(after, txt)
        after = p

    tighten_to_next_heading(doc, heading, next_head_space_before_pt=12)

def write_personal_development(doc: Document, fields: Dict) -> None:
    items = None
    for key in ["professional_development", "personal_development", "development", "training"]:
        if fields.get(key):
            v = fields.get(key)
            if isinstance(v, list):
                items = v
            elif isinstance(v, dict):
                items = []
                for vv in v.values():
                    if isinstance(vv, list): items.extend(vv)
                    elif vv: items.append(str(vv))
            elif v: items = [str(v)]
            break
    _write_simple_bullet_section(doc, "PROFESSIONAL DEVELOPMENT", items or [])

def write_professional_affiliations(doc: Document, fields: Dict) -> None:
    items = None
    for key in ["professional_affiliations", "affiliations", "memberships"]:
        if fields.get(key):
            v = fields.get(key)
            if isinstance(v, list):
                items = v
            elif isinstance(v, dict):
                items = []
                for vv in v.values():
                    if isinstance(vv, list): items.extend(vv)
                    elif vv: items.append(str(vv))
            elif v: items = [str(v)]
            break
    _write_simple_bullet_section(doc, "PROFESSIONAL AFFILIATIONS", items or [])

def write_experience(doc: Document, fields: Dict) -> None:
    heading_aliases = ["EXPERIENCE", "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE", "EMPLOYMENT HISTORY"]
    h_idx = None; chosen = None
    for cand in heading_aliases:
        h_idx = find_heading_index(doc, cand)
        if h_idx is not None:
            chosen = cand; break
    if h_idx is None:
        return

    apply_heading_style(doc, chosen)
    _strip_placeholders_after_heading(doc, chosen)

    anchor = doc.paragraphs[h_idx]
    after = anchor

    exp = fields.get("experience") or []
    def item_key(e): return _year_key(e.get("dates",""))
    exp_sorted = sorted(exp, key=item_key, reverse=True)

    last = after
    for idx_e, e in enumerate(exp_sorted):
        dates = (e.get("dates") or "").strip()
        company = (e.get("company") or "").strip()
        jt_raw = (e.get("job_title") or e.get("title") or "").strip()
        desc = (e.get("description") or "").strip()

        jt_lines = [s for s in re.split(r"\r?\n", jt_raw) if s.strip()]
        title = jt_lines[0] if jt_lines else ""
        intro = " ".join(jt_lines[1:]).strip()

        # single blank between entries (none before the first)
        if idx_e > 0:
            gap = insert_paragraph_after(last, "")
            format_paragraph(gap, font_name="Calibri", font_size_pt=10, space_after_pt=0)
            last = gap

        # Line 1: dates (bold) + company (bold uppercase)
        p1 = insert_paragraph_after(last, "")
        r1 = p1.add_run(dates + "\t\t\t"); set_run(r1, name="Calibri", size_pt=10, bold=True)
        r2 = p1.add_run(company.upper()); set_run(r2, name="Calibri", size_pt=10, bold=True, italic=False)
        last = p1

        # Line 2: job title (bold)
        if title:
            p2 = insert_paragraph_after(p1, "\t\t\t" + title)
            for r in p2.runs: set_run(r, name="Calibri", size_pt=10, bold=True, italic=False)
            last = p2

        # Line 3: intro/summary (italic)
        if intro:
            p3 = insert_paragraph_after(last, "\t\t\t" + intro)
            for r in p3.runs: set_run(r, name="Calibri", size_pt=10, bold=False, italic=True)
            last = p3

        # Bullets
        if desc:
            parts = re.split(r"\n{2,}|;", desc)
            for part in parts:
                part = re.sub(r"-\s*[\r\n]+\s*", "", part)  # un-hyphenate
                part = re.sub(r"[\r\n]+", " ", part)        # merge soft wraps
                part = re.sub(r"\s+", " ", part).strip()
                if part:
                    bp = add_clean_bullet(last, part)
                    last = bp

    tighten_to_next_heading(doc, chosen, next_head_space_before_pt=12)


def _remove_x_placeholders(doc: "Document"):
    """
    Remove any paragraph that is essentially an 'Xxxxxx...' placeholder.
    Criteria: >= 60% of non-space characters are X/x or line is only X/x and punctuation.
    """
    i = 0
    while i < len(doc.paragraphs):
        txt = (doc.paragraphs[i].text or "").strip()
        if not txt:
            i += 1
            continue
        # collapse to just X/x and count
        letters = re.sub(r"\s+", "", txt)
        if letters:
            x_only = re.sub(r"[^xX]", "", letters)
            # if 60%+ are X/x OR it's made up of only X/x and punctuation
            if (len(x_only) >= int(0.6 * len(letters))) or re.fullmatch(r"[xX\s\.,;:!\-–—\(\)]+", txt):
                # delete and continue without advancing
                delete_paragraph(doc.paragraphs[i])
                continue
        i += 1


def _remove_section_if_empty(doc: "Document", heading: str, has_content: bool):
    """
    Remove the section 'heading' and its following body until the next ALL-CAPS heading
    if has_content is False.
    """
    if has_content:
        return
    h_idx = find_heading_index(doc, heading)
    if h_idx is None:
        return
    # delete heading
    delete_paragraph(doc.paragraphs[h_idx])
    # delete body until next heading
    while h_idx < len(doc.paragraphs):
        if h_idx >= len(doc.paragraphs):
            break
        t = (doc.paragraphs[h_idx].text or "").strip()
        if is_heading_text(t):
            break
        delete_paragraph(doc.paragraphs[h_idx])

# ============== Render ==============
def ensure_body_font(doc: Document):
    for p in doc.paragraphs:
        txt = p.text.strip()
        if (txt.isupper() and len(txt) <= 32) or txt.startswith("CURRICULUM VITAE FOR") or txt.startswith("CANDIDATE LOCATION"):
            continue
        for r in p.runs:
            if r.text.strip():
                r.font.name = "Calibri"
                r.font.size = Pt(10)

def render(doc: Document, fields: Dict) -> None:
    for hd in ["PERSONAL PROFILE", "KEY SKILLS", "EDUCATION", "QUALIFICATIONS",
               "PERSONAL DEVELOPMENT", "PROFESSIONAL AFFILIATIONS",
               "EMPLOYMENT HISTORY", "EXPERIENCE", "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE",
               "OTHER HEADINGS", "ADDITIONAL INFORMATION"]:
        apply_heading_style(doc, hd)

    name = (fields.get("name") or "").strip()
    location = (fields.get("location") or "").strip()
    style_first_two_lines(doc, name, location)

    write_summary(doc, fields)
    write_education(doc, fields.get("education") or [])
    write_skills(doc, fields)
    write_personal_development(doc, fields)
    write_professional_affiliations(doc, fields)
    write_qualifications(doc, fields.get("qualifications") or [])
    write_experience(doc, fields)

    _remove_template_placeholders(doc)
    _remove_x_placeholders(doc)
    _remove_section_if_empty(doc, "ADDITIONAL INFORMATION", bool(fields.get("additional_information")))
    _collapse_multiple_blank_paragraphs(doc)
    ensure_body_font(doc)

# ============== CLI ==============
def main():
    ap = argparse.ArgumentParser(description="CV Reformatter — consolidated stable build")
    ap.add_argument("-f", "--fields", default="output/fields.json", help="Path to fields.json")
    ap.add_argument("-t", "--template", default=None, help="Path to .docx template")
    ap.add_argument("-o", "--output", default="output/Reformatted_CV1.docx", help="Output .docx")
    args = ap.parse_args()

    fields = load_json(args.fields)
    template = find_template(args.template)
    doc = Document(template)
    render(doc, fields)
    os.makedirs(os.path.dirname(args.output) or ".", exist_ok=True)
    doc.save(args.output)
    print(f"✔ Wrote CV to: {args.output}")

if __name__ == "__main__":
    main()
